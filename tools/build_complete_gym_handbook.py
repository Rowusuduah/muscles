from __future__ import annotations

import math
import os
import re
import shutil
from pathlib import Path

from PIL import Image, ImageDraw, ImageEnhance, ImageFilter, ImageFont, ImageOps
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"E:\gYM")
SOURCE_DIR = ROOT / "GYM"
BUILD_DIR = Path(os.environ.get("TEMP", str(ROOT))) / "gym_handbook_final_build"
ASSET_DIR = BUILD_DIR / "assets"
OUTPUT_DOCX = ROOT / "Complete_Gym_Equipment_Handbook_Revised.docx"

PAGE_W = Inches(8.5)
PAGE_H = Inches(11)
MARGIN = Inches(1)
CONTENT_W_IN = 6.5
CONTENT_W_DXA = 9360

INK = "17212B"
MUTED = "5F6B76"
LIGHT = "F3F6F8"
RULE = "D5DCE2"
WHITE = "FFFFFF"
GOLD = "D29A2E"
RED = "B84242"
GREEN = "2E7D56"

CATEGORY_COLORS = {
    "Push": "C74B50",
    "Pull": "2E6FA7",
    "Legs": "2E8555",
    "Core": "7356A5",
    "Full Body": "C67A24",
    "Cardio": "16889E",
}


def photo_name(n: int) -> str:
    return "Gym equipment.jpeg" if n == 0 else f"Gym equipment {n}.jpeg"


def guide(
    no: int,
    category: str,
    title: str,
    photos: list[int],
    *,
    confidence: str,
    evidence: str,
    pattern: str,
    level: str,
    role: str,
    purpose: str,
    primary: str,
    secondary: str,
    setup: list[str],
    execution: list[str],
    cues: list[str],
    mistakes: list[tuple[str, str]],
    safety: str,
    alternatives: list[str],
    placement: str,
    programming: str,
    progression: str,
    tempo: str = "2 seconds down, brief control, 1-2 seconds up",
    movement: str = "controlled path",
    annotations: list[str] | None = None,
) -> dict:
    return {
        "no": no,
        "category": category,
        "title": title,
        "photos": photos,
        "confidence": confidence,
        "evidence": evidence,
        "pattern": pattern,
        "level": level,
        "role": role,
        "purpose": purpose,
        "primary": primary,
        "secondary": secondary,
        "setup": setup,
        "execution": execution,
        "cues": cues,
        "mistakes": mistakes,
        "safety": safety,
        "alternatives": alternatives,
        "placement": placement,
        "programming": programming,
        "progression": progression,
        "tempo": tempo,
        "movement": movement,
        "annotations": annotations or ["working handles", "body support", "load or selector", "adjustment point"],
    }


GUIDES = [
    guide(1, "Push", "Plate-Loaded Chest Press", [0], confidence="High", evidence="Matrix branding, seated back support, horizontal converging press arms and plate horns are clearly visible.", pattern="Horizontal push", level="Beginner to advanced", role="Primary compound press", purpose="Build chest pressing strength with a stable torso and independent machine arms.", primary="Pectoralis major", secondary="Triceps, anterior deltoids and serratus anterior", setup=["Load both plate horns evenly and keep the smallest plates closest to the frame.", "Set the seat so the working handles line up near mid-chest and the feet stay flat.", "Sit tall with the head and upper back supported; gently draw the shoulder blades back and down."], execution=["Grip with neutral wrists and elbows slightly below shoulder height.", "Press forward and slightly inward without lifting the shoulders.", "Stop just short of forcefully locking the elbows.", "Return until the chest is comfortably stretched while the back remains on the pad."], cues=["Press the handles away from the chest", "Keep ribs stacked over the pelvis", "Make the return quieter than the press"], mistakes=[("Seat too low", "Raises the elbows and can crowd the shoulders; raise the seat."), ("Bouncing from the bottom", "Use a smaller load and pause under control."), ("Uneven plates", "Match both sides before every set.")], safety="Use the machine stops and never place hands between moving arms and the frame.", alternatives=["Dumbbell bench press", "Barbell bench press", "Cable chest press"], placement="Early in a push or full-body session after the warm-up.", programming="3-4 sets of 6-12 reps, 90-180 seconds rest, usually 1-3 repetitions in reserve (RIR).", progression="When every set reaches the top of the range with identical form, add the smallest equal plate increase.", movement="forward press"),
    guide(2, "Push", "Plate-Loaded Incline Press", [14], confidence="High", evidence="The machine placard reads Incline Press and the angled backrest and upward press path are visible.", pattern="Upward horizontal push", level="Beginner to advanced", role="Primary or secondary compound press", purpose="Emphasize the clavicular chest and anterior deltoids through an inclined pressing path.", primary="Upper pectoralis major", secondary="Anterior deltoids and triceps", setup=["Load each side evenly.", "Adjust the seat so the handles start between upper-chest and shoulder height.", "Keep feet planted and shoulder blades supported against the inclined pad."], execution=["Unweight the handles smoothly from the stops.", "Press up and forward while keeping wrists stacked over forearms.", "Finish with elbows nearly straight, not aggressively locked.", "Lower until the elbows reach a comfortable position slightly behind the torso."], cues=["Drive through the palms", "Keep the sternum tall without over-arching", "Elbows travel under the handles"], mistakes=[("Seat too high", "Turns the exercise into a lower chest press; lower the seat."), ("Flaring elbows", "Use a grip and elbow angle that keep forearms vertical."), ("Losing back contact", "Reduce load and keep the ribs down.")], safety="Do not force a deep bottom position if the front of the shoulder feels pinched.", alternatives=["Incline dumbbell press", "Incline barbell bench", "Low-to-high cable press"], placement="First or second press on a push day.", programming="3-4 sets of 6-12 reps, 90-180 seconds rest, 1-3 RIR.", progression="Add repetitions first, then a small equal plate increase.", movement="upward press"),
    guide(3, "Push", "Plate-Loaded Decline Press", [4], confidence="High", evidence="The seated backrest, downward-forward arm path and plate-loaded press geometry identify a decline press; this is not a pulldown.", pattern="Downward horizontal push", level="Beginner to advanced", role="Compound chest press", purpose="Train the chest and triceps through a slightly downward pressing path with full back support.", primary="Pectoralis major", secondary="Triceps and anterior deltoids", setup=["Match the plates on both sides.", "Set the seat so the handles begin around the lower-to-mid chest.", "Brace the feet and keep the pelvis and upper back on the pad."], execution=["Take the handles with neutral wrists.", "Press forward and slightly down without shrugging.", "Pause before elbow lockout.", "Return slowly to a comfortable chest stretch."], cues=["Press down the track, not toward the ceiling", "Shoulders stay away from ears", "Control both arms evenly"], mistakes=[("Calling it a pulldown", "The machine uses a pressing path and backrest; follow press setup."), ("Over-arching", "Keep ribs down and reduce load."), ("Short return", "Use the largest pain-free range while maintaining shoulder position.")], safety="Confirm the press arms are fully on their stops before entering or leaving.", alternatives=["Slight-decline dumbbell press", "Parallel-bar dip with assistance", "Flat machine chest press"], placement="After a primary chest press or as the main press for a lower-chest emphasis.", programming="3 sets of 8-12 reps, 90-150 seconds rest, 1-3 RIR.", progression="Add one rep per set before adding equal plates.", movement="forward-down press"),
    guide(4, "Push", "Plate-Loaded Shoulder Press", [5, 8], confidence="High", evidence="The visible placard reads Shoulder Press; both photographs show the same seat, back pad and overhead press arms from alternate angles.", pattern="Vertical push", level="Beginner to advanced", role="Compound shoulder press", purpose="Develop overhead pressing strength with a guided path and supported torso.", primary="Anterior and lateral deltoids", secondary="Triceps, upper chest and upper trapezius", setup=["Load both horns evenly.", "Adjust the seat so the working handles start near shoulder height.", "Choose the grip that keeps wrists neutral and forearms close to vertical."], execution=["Brace the trunk against the pad.", "Press upward without pushing the head forward.", "Stop before forceful elbow lockout.", "Lower until the handles return near shoulder level with no pinch."], cues=["Ribs down", "Reach up without shrugging", "Forearms follow the handles"], mistakes=[("Seat too low", "Creates an excessively deep start; raise the seat."), ("Back arch", "Reduce weight and keep the trunk against the pad."), ("Hard lockout", "Finish smoothly with soft elbows.")], safety="Avoid an unusually deep start or painful shoulder rotation.", alternatives=["Dumbbell shoulder press", "Landmine press", "Cable shoulder press"], placement="Early in push training, usually after or instead of a chest press.", programming="3-4 sets of 6-12 reps, 90-180 seconds rest, 1-3 RIR.", progression="Reach the rep target with both arms moving together, then add equal load.", movement="overhead press"),
    guide(5, "Push", "Plate-Loaded Seated Dip", [16], confidence="High", evidence="The plate horns, seated torso pad and downward press handles identify a plate-loaded dip machine.", pattern="Downward push", level="Beginner to advanced", role="Compound triceps accessory", purpose="Train dip mechanics with a supported body and adjustable external load.", primary="Triceps brachii", secondary="Lower chest and anterior deltoids", setup=["Load the plate horns evenly.", "Set the seat so the handles sit beside the lower ribs.", "Choose the handle position that keeps wrists neutral and shoulders comfortable."], execution=["Sit tall and depress the shoulders.", "Press the handles down until the elbows are nearly straight.", "Pause without leaning onto the mechanism.", "Return until the elbows bend comfortably without the shoulders rolling forward."], cues=["Push the shoulders away from the ears", "Elbows travel behind and then down", "Keep the torso still"], mistakes=[("Shrugging", "Reduce load and reset shoulder depression."), ("Excessive forward lean", "Use only a small deliberate lean if emphasizing chest."), ("Half repetitions", "Use a controlled pain-free elbow range.")], safety="Keep fingers clear of plate horns, stops and lever pivots.", alternatives=["Assisted dip", "Cable pressdown", "Close-grip bench press"], placement="After primary presses or as a triceps-focused compound movement.", programming="2-4 sets of 8-15 reps, 75-150 seconds rest, 1-3 RIR.", progression="Add repetitions, then the smallest even plate increase.", movement="downward press"),
    guide(6, "Push", "Assisted Pull-Up / Dip Machine", [35], confidence="High", evidence="Kneeling assistance pad, pull-up handles, dip handles and a selectorized counterweight stack are visible.", pattern="Vertical pull or downward push", level="Beginner to intermediate", role="Bodyweight skill builder", purpose="Reduce effective bodyweight so pull-ups, chin-ups and dips can be practiced with consistent technique.", primary="Pull-ups: lats and upper back; dips: triceps and chest", secondary="Biceps, shoulders, forearms and core", setup=["Select assistance before climbing on; more selected weight means more help.", "Use the steps and hold a fixed handle while placing both knees on the pad.", "Choose pull-up or dip handles and center the body."], execution=["For pull-ups, begin long through the arms, drive elbows down and bring the upper chest toward the handles.", "For dips, keep shoulders depressed and lower only as far as comfortable.", "Move without swinging or bouncing the assistance pad.", "Step off only after the pad returns to its start position."], cues=["More stack weight equals easier", "Move the body, not the neck", "Control the pad both ways"], mistakes=[("Selecting too little assistance", "Use enough help for smooth full repetitions."), ("Kipping", "Slow the tempo and keep legs quiet."), ("Stepping off a raised pad", "Return the mechanism fully first.")], safety="The pad moves upward quickly when unloaded; maintain contact until it settles.", alternatives=["Lat pulldown", "Band-assisted pull-up", "Plate-loaded seated dip"], placement="Early in pull training for pull-ups or after presses for dips.", programming="3-4 sets of 5-10 controlled reps, 90-180 seconds rest, 1-3 RIR.", progression="Reduce assistance by one small increment only after all repetitions are controlled.", movement="assisted bodyweight path"),
    guide(7, "Push", "Pec Fly / Rear-Delt Machines", [34, 46], confidence="High", evidence="Two selectorized fly stations are shown; their adjustable arms and exercise placards support chest-fly and reverse-fly use.", pattern="Horizontal adduction or abduction", level="Beginner to advanced", role="Isolation accessory", purpose="Isolate chest adduction when facing away from the pad and rear-shoulder abduction when facing the pad.", primary="Chest-fly: pectoralis major; reverse-fly: posterior deltoids", secondary="Front deltoids or rhomboids and middle trapezius", setup=["Select the function and set the start-angle lever symmetrically.", "Adjust seat height so the handles align near shoulder level.", "Choose a light load first because the long lever increases difficulty."], execution=["Chest fly: keep a soft elbow and sweep the arms together as if hugging a barrel.", "Rear-delt fly: face the pad and open the arms without shrugging.", "Pause briefly at the shortened position.", "Return slowly without letting the stack slam."], cues=["Move from the shoulder, not the hand", "Soft elbows", "Keep the chest on the pad for rear delts"], mistakes=[("Overstretching", "Use a smaller start angle."), ("Bending and straightening elbows", "Freeze a soft elbow angle."), ("Shrugging", "Reduce load and keep shoulders down.")], safety="Set both adjustable arms before sitting and keep fingers away from arm-position gears.", alternatives=["Cable fly", "Dumbbell reverse fly", "Face pull"], placement="After compound presses for chest fly or after rows for rear delts.", programming="2-4 sets of 10-20 reps, 45-90 seconds rest, 1-3 RIR.", progression="Add repetitions with a quiet weight stack before increasing one pin.", movement="arms sweep in or out"),
    guide(8, "Push", "Flat Barbell Bench Press Station", [31], confidence="High", evidence="A flat bench, bar catches and Olympic barbell station are clearly visible.", pattern="Horizontal push", level="Intermediate skill demand", role="Primary free-weight press", purpose="Build general chest and pressing strength with a free barbell.", primary="Pectoralis major", secondary="Triceps and anterior deltoids", setup=["Set the bar so the elbows remain slightly bent when unracking.", "Center the bench, place eyes under the bar and secure plates with collars.", "Use a spotter or correctly positioned safety arms for challenging sets."], execution=["Plant the feet and retract the shoulder blades.", "Unrack over the shoulder joints.", "Lower under control toward the mid-to-lower chest.", "Press back over the shoulders without bouncing or lifting the hips."], cues=["Feet drive into the floor", "Touch softly", "Wrists stacked over elbows"], mistakes=[("Bar too high in rack", "Lower the hooks so unracking does not protract the shoulders."), ("Bouncing", "Reduce load and pause lightly."), ("No safety plan", "Use a spotter or safeties.")], safety="Never use collars without a spotter when no safety arms are available; facility rules may differ.", alternatives=["Machine chest press", "Dumbbell bench press", "Push-up"], placement="First strength movement on a push or upper-body day.", programming="3-5 sets of 4-10 reps, 2-4 minutes rest, generally 1-3 RIR.", progression="Increase load only when the same bar path and touch point are repeatable.", movement="bar lowers and presses"),
    guide(9, "Push", "Incline Barbell Bench Press Station", [32], confidence="High", evidence="The bench and bar catches are fixed at an incline suitable for barbell pressing.", pattern="Inclined push", level="Intermediate skill demand", role="Primary or secondary press", purpose="Train upper-chest and shoulder pressing strength with a free bar.", primary="Upper pectoralis major", secondary="Anterior deltoids and triceps", setup=["Position eyes just behind the bar and feet securely on the floor or platform.", "Set bar height for a short, safe unrack.", "Use collars, a spotter or safety equipment according to facility rules."], execution=["Unrack to the shoulder line.", "Lower toward the upper chest with forearms nearly vertical.", "Pause softly without bouncing.", "Press up and slightly back to the shoulder line."], cues=["Upper back stays pinned", "Touch high on the chest", "Press back toward the rack"], mistakes=[("Bench angle too steep for the goal", "Use the fixed station as designed and expect more shoulder work."), ("Elbows flared to 90 degrees", "Keep elbows slightly tucked."), ("Hips lifting", "Reduce load and improve foot position.")], safety="Use a competent spotter for hard sets; do not attempt to roll a trapped bar down the torso.", alternatives=["Plate-loaded incline press", "Incline dumbbell press", "Low-to-high cable press"], placement="First or second exercise on push or upper-body days.", programming="3-4 sets of 5-10 reps, 2-3 minutes rest, 1-3 RIR.", progression="Add small increments while preserving touch point and back contact.", movement="inclined bar press"),
    guide(10, "Push", "Decline Barbell Bench Press Station", [33], confidence="High", evidence="The downward bench angle and leg-retention rollers identify a decline bench-press station.", pattern="Declined push", level="Intermediate skill demand", role="Secondary compound press", purpose="Press a barbell on a decline while the leg rollers secure the body.", primary="Pectoralis major", secondary="Triceps and anterior deltoids", setup=["Secure both legs under the rollers before handling the bar.", "Check rack height, collars and spotter position.", "Retract shoulder blades and keep the head supported."], execution=["Unrack over the lower chest.", "Lower to a consistent lower-chest touch point.", "Press smoothly without bouncing.", "Re-rack only when both hooks are clearly engaged."], cues=["Stay anchored under the rollers", "Touch low and controlled", "Confirm both rack hooks"], mistakes=[("Sliding on the bench", "Reposition fully under the leg rollers."), ("Bar drifting toward the face", "Press toward the rack only after clearing the chest."), ("Solo maximal attempts", "Use a spotter.")], safety="Getting in and out is less convenient than a flat bench; unload the bar before changing the setup.", alternatives=["Plate-loaded decline press", "Flat bench press", "Assisted dip"], placement="After a main press or as an optional chest variation.", programming="3 sets of 6-12 reps, 2-3 minutes rest, 1-3 RIR.", progression="Use small load increases and stable leg anchoring.", movement="declined bar press"),

    guide(11, "Pull", "Plate-Loaded Lat Pulldown", [3], confidence="High", evidence="High independent handles, thigh rollers, seat and plate horns identify a plate-loaded pulldown; this is the image omitted from the draft pulldown page.", pattern="Vertical pull", level="Beginner to advanced", role="Primary back movement", purpose="Develop lats and upper-back strength with independent high handles.", primary="Latissimus dorsi and teres major", secondary="Biceps, brachialis, rhomboids and lower trapezius", setup=["Load both horns evenly.", "Adjust the thigh restraint so the hips stay on the seat.", "Use the steps or handles to reach the high grips safely."], execution=["Begin with arms long and chest gently lifted.", "Drive elbows down toward the ribs.", "Pause when the handles approach shoulder or upper-chest level.", "Return to a controlled overhead stretch without lifting off the seat."], cues=["Elbows into back pockets", "Keep the ribs stacked", "Let the shoulder blades rotate upward on the return"], mistakes=[("Using image 4 for this guide", "Image 3 is the actual plate-loaded pulldown."), ("Pulling behind the neck", "Pull in front to a comfortable upper-chest path."), ("Rocking backward", "Reduce load and keep torso movement small.")], safety="Do not climb on plates or lever arms; use the designed seat and access points.", alternatives=["Selectorized lat pulldown", "Assisted pull-up", "Single-arm cable pulldown"], placement="Early in pull or full-body training.", programming="3-4 sets of 6-12 reps, 90-180 seconds rest, 1-3 RIR.", progression="Add repetitions symmetrically, then equal plates.", movement="handles pull downward"),
    guide(12, "Pull", "Selectorized Lat Pulldown", [42], confidence="High", evidence="Life Fitness cable tower, overhead bar, seat and thigh restraint are visible.", pattern="Vertical pull", level="Beginner to advanced", role="Primary or secondary back movement", purpose="Provide easily adjustable vertical pulling for back development and pull-up preparation.", primary="Latissimus dorsi", secondary="Biceps, brachialis, rhomboids and lower trapezius", setup=["Insert the selector pin fully and add any small increment before sitting.", "Adjust the thigh pad snugly over the thighs.", "Choose a grip that permits neutral wrists and a comfortable shoulder width."], execution=["Sit tall with arms long.", "Pull the bar toward the upper chest by driving elbows down.", "Pause without leaning far backward.", "Return until the arms are long and the stack remains controlled."], cues=["Chest tall, ribs quiet", "Pull elbows, not hands", "No stack slam"], mistakes=[("Pulling behind the neck", "Use an in-front path."), ("Excessive torso swing", "Reduce load."), ("Partial stretch", "Return until shoulders are comfortable and arms long.")], safety="Check that attachments and carabiners are fully closed before pulling.", alternatives=["Plate-loaded pulldown", "Assisted pull-up", "Kneeling cable pulldown"], placement="Early in pull sessions or as a moderate-load back accessory.", programming="3-4 sets of 8-15 reps, 75-150 seconds rest, 1-3 RIR.", progression="Add one pin only when the stack stays quiet and torso remains controlled.", movement="bar pulls down"),
    guide(13, "Pull", "Seated Cable Row", [41], confidence="High", evidence="Low cable, footplates, bench and rowing attachment identify a seated low-row station.", pattern="Horizontal pull", level="Beginner to advanced", role="Compound back movement", purpose="Develop back thickness while the cable maintains tension through the range.", primary="Rhomboids, middle trapezius and lats", secondary="Rear deltoids, biceps, forearms and spinal stabilizers", setup=["Select the load and secure the attachment.", "Place feet on the plates with knees softly bent.", "Sit tall and begin with arms long without rounding aggressively."], execution=["Initiate by drawing the shoulder blades back and pulling elbows toward the lower ribs.", "Pause with the torso nearly vertical.", "Extend the arms first, then allow a controlled shoulder-blade reach.", "Keep the cable aligned with the attachment."], cues=["Elbows to back pockets", "Long spine", "Reach without collapsing"], mistakes=[("Turning it into a torso swing", "Reduce load and keep hip motion small."), ("Shrugging", "Pull toward lower ribs."), ("Curling with arms only", "Initiate with the back.")], safety="Keep feet on the plates and never let a loose handle snap toward the stack.", alternatives=["Chest-supported row", "One-arm dumbbell row", "Machine row"], placement="First or second pull movement.", programming="3-4 sets of 8-15 reps, 75-150 seconds rest, 1-3 RIR.", progression="Increase load after full controlled reaches and consistent pauses.", movement="handle rows toward torso"),
    guide(14, "Pull", "Chest-Supported Plate Row", [2], confidence="High", evidence="Angled chest pad, seat, low independent handles and plate horns identify a supported row.", pattern="Horizontal pull", level="Beginner to advanced", role="Primary or secondary back movement", purpose="Train horizontal pulling while the chest pad limits lower-back fatigue and torso momentum.", primary="Lats, rhomboids and middle trapezius", secondary="Rear deltoids, biceps and forearms", setup=["Load the horns evenly.", "Adjust the seat so the chest is supported and handles are reachable with long arms.", "Choose neutral or overhand grips based on comfort."], execution=["Keep the chest on the pad.", "Pull elbows behind the torso without shrugging.", "Pause briefly at the top.", "Lower until arms are long and shoulder blades can move forward naturally."], cues=["Glue chest to pad", "Pull elbows, not shoulders", "Control both arms"], mistakes=[("Leaving the chest pad", "Reduce load."), ("Shortening the return", "Allow a controlled reach."), ("Uneven loading", "Match plates.")], safety="Do not place hands near the moving lever pivots while loading or unloading.", alternatives=["Seated cable row", "Dumbbell chest-supported row", "FreeMotion row"], placement="Early on pull day or after a vertical pull.", programming="3-4 sets of 6-12 reps, 90-150 seconds rest, 1-3 RIR.", progression="Add repetitions first, then equal plate increments.", movement="elbows row backward"),
    guide(15, "Pull", "Independent Plate-Loaded Row", [28], confidence="Medium-high", evidence="Chest pad, seat, independent lever arms and plate horns identify a plate-loaded rowing station; exact model is not legible.", pattern="Horizontal pull", level="Beginner to advanced", role="Back compound/accessory", purpose="Allow unilateral or bilateral rows with chest support and independent arms.", primary="Lats and upper back", secondary="Rear deltoids, biceps and grip", setup=["Load each side for bilateral work or one side for supported unilateral work.", "Set seat height so the chest pad supports the sternum without crowding the neck.", "Grip the selected handles with neutral wrists."], execution=["Brace against the pad.", "Row one or both handles toward the ribs.", "Pause without twisting.", "Lower to a full controlled reach."], cues=["Stay square", "Elbow follows the handle", "Reach long, then row"], mistakes=[("Rotating during bilateral reps", "Reduce load and keep both shoulders level."), ("Pulling too high", "Aim toward the rib cage."), ("Dropping plates", "Guide the arms to the stops.")], safety="For unilateral sets, stabilize with the free hand only on a fixed frame member.", alternatives=["Chest-supported plate row", "Seated cable row", "One-arm cable row"], placement="Second or third movement on a pull day.", programming="3 sets of 8-15 reps per side, 60-120 seconds rest, 1-3 RIR.", progression="Add reps until both sides match, then increase load conservatively.", movement="independent row"),
    guide(16, "Pull", "Plate-Loaded Biceps Curl", [10], confidence="High", evidence="Arm pad, curling handles, pivot and plate horns identify a plate-loaded biceps curl.", pattern="Elbow flexion", level="Beginner to advanced", role="Isolation accessory", purpose="Stabilize the upper arms while loading elbow flexion.", primary="Biceps brachii", secondary="Brachialis, brachioradialis and forearms", setup=["Load both sides evenly if the machine has independent horns.", "Adjust the seat so elbows align near the machine pivot and upper arms rest on the pad.", "Use the grip that keeps wrists comfortable."], execution=["Begin with elbows nearly straight but not forced.", "Curl without lifting the upper arms.", "Squeeze briefly at the top.", "Lower fully under control."], cues=["Upper arms stay heavy on pad", "Knuckles follow the handle", "Slow bottom"], mistakes=[("Elbows lifting", "Lower the load and keep contact with the pad."), ("Wrist bending", "Use a neutral stacked wrist."), ("Dropping the bottom", "Control the eccentric.")], safety="Do not hyperextend the elbows at the bottom or load one side accidentally.", alternatives=["Selectorized biceps curl", "Preacher curl", "Cable curl"], placement="After back compounds.", programming="2-4 sets of 8-15 reps, 60-90 seconds rest, 1-3 RIR.", progression="Add reps with a controlled bottom before adding plates.", movement="handles curl upward"),
    guide(17, "Pull", "Selectorized Biceps Curl", [36, 37], confidence="High", evidence="The same Life Fitness biceps-curl station is shown from front and side; the exercise placard and arm pad are visible.", pattern="Elbow flexion", level="Beginner to advanced", role="Isolation accessory", purpose="Provide quick load changes and a stable upper-arm position for biceps training.", primary="Biceps brachii", secondary="Brachialis and forearms", setup=["Insert the weight pin completely.", "Adjust the seat so elbows align with the pivot and the upper arms rest flat.", "Choose the handle setting that keeps wrists neutral."], execution=["Start with arms long and shoulders relaxed.", "Curl while keeping elbows on the pad.", "Pause without rolling shoulders forward.", "Lower until the stack is nearly down, then reverse smoothly."], cues=["Elbows stay planted", "Squeeze without shrugging", "Quiet stack"], mistakes=[("Seat too high", "Elbows sit above the pivot; lower the seat."), ("Shoulder roll", "Reduce load."), ("Stack impact", "Slow the final third of the return.")], safety="Keep fingers away from the selector stack and confirm the pin is fully seated.", alternatives=["Plate-loaded biceps curl", "Preacher curl", "Dumbbell curl"], placement="Late in pull or upper-body sessions.", programming="2-4 sets of 10-20 reps, 45-90 seconds rest, 1-3 RIR.", progression="Add reps first, then one pin increment.", movement="handles curl upward"),
    guide(18, "Pull", "Preacher Curl Benches", [26, 27], confidence="High", evidence="Both photographs show angled arm pads and bar rests designed for preacher curls.", pattern="Elbow flexion", level="Beginner to intermediate", role="Free-weight biceps accessory", purpose="Limit torso swing by supporting the upper arms on an angled pad.", primary="Biceps brachii", secondary="Brachialis and forearms", setup=["Adjust the seat so the armpits sit near the top edge without rounding forward.", "Select an EZ bar or dumbbells that can be controlled from the bottom.", "Use a spotter or safe handoff for heavy bars."], execution=["Keep upper arms in contact with the pad.", "Curl until forearms approach vertical.", "Pause without moving the shoulders.", "Lower to near full elbow extension without forcefully locking."], cues=["Pad supports the arm", "Shoulders stay behind elbows", "Own the last few inches down"], mistakes=[("Seat too low", "Forces shoulders to roll; raise the seat."), ("Relaxing at bottom", "Maintain tension and a soft elbow."), ("Using momentum", "Reduce load.")], safety="The stretched bottom is demanding; avoid sudden extension or dropping a heavy bar into the rests.", alternatives=["Selectorized curl", "Cable preacher curl", "Incline dumbbell curl"], placement="After rows and pulldowns.", programming="2-4 sets of 8-15 reps, 60-120 seconds rest, 1-3 RIR.", progression="Add reps before choosing the next fixed bar or adding plates.", movement="forearms curl toward shoulders"),
    guide(19, "Pull", "Fixed EZ-Curl Bars", [21], confidence="High", evidence="A rack of labeled fixed-weight EZ curl bars is clearly visible.", pattern="Elbow flexion or extension", level="Beginner to advanced", role="Free-weight accessory tool", purpose="Provide fast, balanced weight selection for curls, reverse curls and controlled triceps extensions.", primary="Exercise-dependent: biceps or triceps", secondary="Forearms, brachialis and grip", setup=["Read the labeled total weight before lifting.", "Stand close and lift with both hands using the legs rather than reaching from the back.", "Choose the inner or outer angled grip that keeps wrists comfortable."], execution=["For curls, keep elbows near the torso and raise the bar without leaning back.", "For reverse curls, use a lighter weight and overhand grip.", "For extensions, maintain control above the head or on a bench.", "Return the bar to the matching labeled slot."], cues=["Bar stays level", "Wrists follow the angles", "No torso swing"], mistakes=[("Assuming each end label is per side", "Treat the rack label as total bar weight unless the facility says otherwise."), ("Wide swinging arc", "Keep elbows controlled."), ("Leaving bars on floor", "Re-rack immediately.")], safety="Inspect fixed collars and do not use a bar with a loose end or damaged grip.", alternatives=["Dumbbells", "Cable curls", "Plate-loaded curl machine"], placement="Late in pull or arm training.", programming="2-4 sets of 8-20 reps, 45-90 seconds rest, 1-3 RIR.", progression="Move to the next bar only after the current weight reaches the top of the rep range.", movement="bar curls or extends"),
    guide(20, "Pull", "Power Tower", [43], confidence="High", evidence="Pull-up bars, dip handles and vertical-knee-raise elbow pads are visible on one station.", pattern="Vertical pull, downward push and trunk flexion", level="Intermediate; scalable with assistance elsewhere", role="Bodyweight multi-station", purpose="Provide pull-ups, chin-ups, dips and supported knee raises in one footprint.", primary="Exercise-dependent: back, chest, triceps or abdominal muscles", secondary="Grip, shoulders, hip flexors and trunk stabilizers", setup=["Choose the station and confirm the floor and steps are clear.", "Use the steps rather than jumping to the handles.", "For knee raises, place forearms fully on pads and back against support."], execution=["Pull-ups: move from a controlled hang and drive elbows down.", "Dips: depress shoulders and lower only as far as comfortable.", "Knee raises: curl the pelvis rather than swinging the legs.", "Step down under control after the final repetition."], cues=["Long body, quiet legs", "Shoulders away from ears", "Move the pelvis on knee raises"], mistakes=[("Kipping", "Use assistance or fewer reps."), ("Deep painful dips", "Shorten range or use assisted dips."), ("Swinging knee raises", "Pause between reps.")], safety="Do not use if grips, pads or steps are wet; give other users clearance around the station.", alternatives=["Assisted pull-up/dip", "Lat pulldown", "Captain's-chair knee raise"], placement="Early for pull-ups, later for dips or core work.", programming="3-5 sets of 3-10 quality reps; stop before technique deteriorates.", progression="Add reps, reduce assistance on another station, or use slower tempo before external load.", movement="bodyweight vertical path"),
    guide(21, "Pull", "Back Extension / Roman Chair", [23], confidence="High", evidence="Life Fitness frame label, hip pads, ankle rollers and angled body position identify a back-extension bench.", pattern="Hip hinge", level="Beginner to advanced", role="Posterior-chain accessory", purpose="Train hip extension and trunk endurance with the body supported at the hips and ankles.", primary="Erector spinae, gluteus maximus and hamstrings", secondary="Adductors and deep trunk stabilizers", setup=["Adjust the hip pads so their top edge sits below the hip crease.", "Secure ankles beneath the rollers with feet on the platform.", "Begin with bodyweight and arms crossed."], execution=["Hinge forward from the hips while keeping a long spine.", "Descend only as far as hamstring length and pad position allow.", "Drive hips into the pad to return.", "Stop when the torso aligns with the legs; do not hyperextend."], cues=["Fold at hips", "Long neck", "Finish straight, not arched"], mistakes=[("Pad on the abdomen", "Lower it below the hip crease."), ("Rounding and snapping up", "Slow down and use bodyweight."), ("Hyperextending", "Stop in a straight line.")], safety="If the exercise produces sharp back pain, stop and seek a qualified assessment.", alternatives=["Cable pull-through", "Romanian deadlift", "Glute bridge"], placement="After primary lower-body or back work.", programming="2-4 sets of 8-20 reps, 60-120 seconds rest, 2-4 RIR for beginners.", progression="Add reps and tempo before holding a light plate close to the chest.", movement="torso hinges at hips"),

    guide(22, "Legs", "45-Degree Plate-Loaded Leg Press", [11], confidence="High", evidence="Angled sled, large footplate, back pad, plate horns and safety handles identify a 45-degree leg press.", pattern="Knee and hip extension", level="Beginner to advanced", role="Primary lower-body compound", purpose="Load the quadriceps and glutes with back support and less balance demand than a free squat.", primary="Quadriceps and gluteus maximus", secondary="Hamstrings, adductors and calves", setup=["Load both sides evenly and keep plates fully on the horns.", "Adjust the back pad and place feet about shoulder width on the platform.", "Unrack using the safety handles only after both feet and the back are set."], execution=["Lower the sled by bending hips and knees while keeping the whole foot planted.", "Stop before the pelvis rolls away from the pad.", "Press through midfoot and heel.", "Re-engage both safeties before relaxing."], cues=["Knees track with toes", "Back stays on pad", "Push the platform away"], mistakes=[("Depth that lifts the pelvis", "Use a shallower pain-free range."), ("Knees collapsing inward", "Reduce load and track knees over toes."), ("Unlocking before set", "Brace first, then release safeties.")], safety="Confirm both mechanical safeties are engaged before loading, unloading or exiting.", alternatives=["Smith squat", "Goblet squat", "Split squat"], placement="First or second lower-body exercise.", programming="3-5 sets of 6-15 reps, 2-4 minutes rest, 1-3 RIR.", progression="Add reps within the range before equal plate increases.", movement="sled lowers and presses"),
    guide(23, "Legs", "Matrix Glute Trainer", [9], confidence="High", evidence="Matrix plate-loaded design, floating hip pad, angled foot platform and exercise diagram match the official Matrix Glute Trainer.", pattern="Hip extension", level="Beginner to advanced", role="Primary or accessory glute movement", purpose="Train a supported bridge-style hip extension with a moving hip pad and stable foot platform.", primary="Gluteus maximus", secondary="Hamstrings, adductors and trunk stabilizers", setup=["Load the horn and secure the plate fully.", "Sit on the bench and place both feet on the angled platform.", "Rotate or position the hip pad across the hip crease according to the machine diagram."], execution=["Begin with hips flexed and ribs stacked.", "Drive through the whole foot and extend the hips.", "Finish when torso and thighs align without arching the lower back.", "Lower until the mechanism reaches a comfortable start position."], cues=["Belt line moves up", "Ribs down", "Squeeze glutes, not lower back"], mistakes=[("Using it as a leg curl", "This is a glute bridge/hip-extension machine."), ("Feet too far away", "Adjust for near-vertical shins at the top."), ("Hyperextending", "Stop at neutral hips.")], safety="Keep hands on fixed handles and fingers away from the moving hip pad and lever.", alternatives=["Booty Builder hip thrust", "Smith-machine hip thrust", "Glute bridge"], placement="Early or middle of a leg/glute session.", programming="3-4 sets of 8-15 reps, 90-180 seconds rest, 1-3 RIR.", progression="Add reps, then a small plate; maintain a one-second top squeeze.", movement="hips extend upward"),
    guide(24, "Legs", "Booty Builder Hip Thrust", [13], confidence="High", evidence="Booty Builder branding, raised footplate, seat and padded hip belt are clearly visible.", pattern="Hip extension", level="Beginner to advanced", role="Primary glute compound", purpose="Provide a dedicated plate-loaded hip thrust with raised feet and a padded restraint.", primary="Gluteus maximus", secondary="Hamstrings, adductors and core stabilizers", setup=["Load both horns evenly if the model uses dual loading points.", "Sit with shoulder blades supported by the back pad and feet on the raised platform.", "Center and secure the padded belt over the hip crease."], execution=["Lower the hips under control while maintaining belt contact.", "Drive through the whole foot.", "Extend until the trunk and thighs form a straight line.", "Pause, then lower without bouncing off the stops."], cues=["Chin slightly tucked", "Ribs down", "Shins close to vertical at top"], mistakes=[("Belt on abdomen", "Place it across the hip crease."), ("Driving from toes", "Use the whole foot."), ("Lumbar overextension", "Finish with glutes and neutral ribs.")], safety="Latch the restraint exactly as shown on the machine and confirm it is secure before loading effort.", alternatives=["Matrix glute trainer", "Smith hip thrust", "Floor glute bridge"], placement="Early in glute-focused leg sessions.", programming="3-5 sets of 6-15 reps, 90-180 seconds rest, 1-3 RIR.", progression="Add reps and top pauses before adding equal plates.", movement="hips thrust upward"),
    guide(25, "Legs", "Plate-Loaded Lying Leg Curl", [12], confidence="High", evidence="Matrix branding, prone pads, long ankle roller, cam adjustment and plate horn identify a plate-loaded lying leg curl.", pattern="Knee flexion", level="Beginner to advanced", role="Hamstring isolation", purpose="Train knee flexion while lying prone with plate-loaded resistance.", primary="Hamstrings", secondary="Gastrocnemius and popliteus", setup=["Load the plate horn and secure the plate.", "Adjust the roller so it rests just above the heels.", "Set the start angle and align the knees close to the machine pivot."], execution=["Brace hips into the pad and hold the handles.", "Curl heels toward the glutes without lifting the pelvis.", "Pause briefly.", "Lower to near full knee extension without dropping the lever."], cues=["Hips stay heavy", "Curl the roller, not the toes", "Slow final third down"], mistakes=[("Calling image 9 a leg curl", "Image 12 is the plate-loaded curl; image 9 is a glute trainer."), ("Roller on calves", "Move it closer to the ankles."), ("Hips lifting", "Reduce load or range.")], safety="Adjust the cam and roller only when the lever is resting and unloaded.", alternatives=["Selectorized lying leg curl", "Seated leg curl", "Stability-ball leg curl"], placement="After a main leg compound or before hip hinges when hamstrings are the priority.", programming="3-4 sets of 8-15 reps, 60-120 seconds rest, 1-3 RIR.", progression="Add reps with full control, then the smallest plate.", movement="heels curl toward glutes"),
    guide(26, "Legs", "Selectorized Lying Leg Curl", [40], confidence="High", evidence="Life Fitness placard reads Leg Curl and the prone bench with ankle roller and weight stack is visible.", pattern="Knee flexion", level="Beginner to advanced", role="Hamstring isolation", purpose="Provide prone knee-flexion training with quick selectorized loading.", primary="Hamstrings", secondary="Gastrocnemius", setup=["Insert the selector pin fully.", "Lie face down with knees near the pivot and roller above the heels.", "Grip the handles and keep the hips supported."], execution=["Curl the roller toward the glutes.", "Keep the front of the hips on the pad.", "Pause near the top without cramping.", "Return until legs are nearly straight and the stack stays controlled."], cues=["Hips down", "Heels lead", "Quiet stack"], mistakes=[("Pelvis lifting", "Reduce load."), ("Fast eccentric", "Use at least two seconds down."), ("Toes pointed hard", "Use a relaxed or neutral ankle unless deliberately varying calf involvement.")], safety="Do not adjust the roller while the stack is lifted.", alternatives=["Plate-loaded lying curl", "Seated leg curl", "Romanian deadlift"], placement="Middle or late in leg training.", programming="2-4 sets of 10-20 reps, 60-90 seconds rest, 1-3 RIR.", progression="Add repetitions before moving one pin.", movement="ankle roller curls upward"),
    guide(27, "Legs", "Selectorized Seated Leg Curl", [39], confidence="High", evidence="The Life Fitness placard reads Seated Leg Curl; thigh restraint, ankle roller and selectorized stack are visible.", pattern="Knee flexion", level="Beginner to advanced", role="Hamstring isolation", purpose="Train hamstrings in a seated, lengthened-hip position with a thigh restraint.", primary="Hamstrings", secondary="Gastrocnemius", setup=["Adjust the back pad so knees line up with the pivot.", "Place the lower roller above the heels and lock the thigh pad snugly over the legs.", "Select the load and start angle."], execution=["Hold the handles and curl the lower roller down and back.", "Keep thighs under the restraint.", "Pause at the bottom.", "Return slowly to near-straight knees."], cues=["Pin thighs down", "Curl under the seat", "Do not chase range with pelvis movement"], mistakes=[("Loose thigh pad", "Lower it until secure but comfortable."), ("Knees off pivot", "Adjust the back pad."), ("Stack slam", "Slow the return.")], safety="Release the thigh restraint only after the lever and stack are fully down.", alternatives=["Lying leg curl", "Cable leg curl", "Romanian deadlift"], placement="After compounds or early when hamstrings are a priority.", programming="3-4 sets of 8-15 reps, 60-120 seconds rest, 1-3 RIR.", progression="Add reps, then one selector increment.", movement="lower legs curl downward"),
    guide(28, "Legs", "Plate-Loaded Leg Extension", [15], confidence="High", evidence="Seat, knee pivot, ankle roller and plate horn identify a plate-loaded knee-extension machine.", pattern="Knee extension", level="Beginner to advanced", role="Quadriceps isolation", purpose="Load the quadriceps directly through open-chain knee extension.", primary="Quadriceps", secondary="Hip flexors act minimally for stabilization", setup=["Load the horn and secure the plate.", "Adjust the back/seat so knees align with the pivot.", "Place the roller just above the ankles and set the start angle."], execution=["Hold the handles and keep the thighs on the seat.", "Extend the knees until nearly straight.", "Squeeze the quadriceps briefly.", "Lower slowly to the selected bend."], cues=["Knee matches pivot", "Lift the roller with the shins", "Smooth top"], mistakes=[("Pivot misalignment", "Reposition seat before loading."), ("Kicking rapidly", "Reduce load and slow the eccentric."), ("Hips lifting", "Use handles and lower resistance.")], safety="Use a pain-free knee range; discomfort at the kneecap is a reason to stop and reassess setup.", alternatives=["Selectorized leg extension", "Split squat", "Leg press"], placement="After compound leg work or early for controlled pre-fatigue.", programming="2-4 sets of 10-20 reps, 60-90 seconds rest, 1-3 RIR.", progression="Increase reps and pause quality before adding a plate.", movement="shins extend forward"),
    guide(29, "Legs", "Selectorized Leg Extension", [38], confidence="High", evidence="Life Fitness placard reads Leg Extension; back pad, ankle roller and stack are visible.", pattern="Knee extension", level="Beginner to advanced", role="Quadriceps isolation", purpose="Provide quickly adjustable quadriceps training with a guided path.", primary="Quadriceps", secondary="Minimal stabilizer demand", setup=["Adjust the back pad so knees align with the pivot.", "Set the ankle roller above the ankles and choose the start angle.", "Insert the selector pin completely."], execution=["Grip the handles and keep hips down.", "Extend until knees are almost straight.", "Pause and contract the quadriceps.", "Lower until the selected start angle with a quiet stack."], cues=["Lift from knees", "Keep hips heavy", "Smooth lockout"], mistakes=[("Seat too far back", "Knees sit behind the pivot; move forward."), ("Momentum", "Reduce load."), ("Hard lockout", "Finish without snapping the knees.")], safety="Stop for sharp joint pain; muscular effort in the front thigh is expected, joint pain is not.", alternatives=["Plate-loaded leg extension", "Leg press", "Step-up"], placement="Middle or late in leg sessions.", programming="2-4 sets of 10-20 reps, 45-90 seconds rest, 1-3 RIR.", progression="Add reps, then one pin while keeping a controlled top.", movement="lower legs extend"),
    guide(30, "Legs", "Hip Adduction Machine", [44, 45], confidence="High", evidence="Both views show the same Life Fitness placard reading Hip Adduction; no hip-abduction machine is documented in the photo set.", pattern="Hip adduction", level="Beginner to advanced", role="Inner-thigh isolation", purpose="Strengthen the hip adductors by bringing the thighs inward against resistance.", primary="Adductor magnus, longus and brevis", secondary="Gracilis, pectineus and trunk stabilizers", setup=["Select a light load and a comfortable starting width.", "Sit fully against the back pad with inner thighs against the pads.", "Use the range lever only while the mechanism is resting."], execution=["Hold the handles and bring the pads together smoothly.", "Pause without bouncing.", "Allow the legs to open under control.", "Stop before the inner thigh is forced into an uncomfortable stretch."], cues=["Squeeze inward", "Pelvis stays level", "Control the opening"], mistakes=[("Calling one view abduction", "Both images are labeled Hip Adduction."), ("Starting too wide", "Reduce range."), ("Stack slamming", "Use slower return.")], safety="Avoid forcing extreme starting width; adductor strains often occur when range exceeds control.", alternatives=["Cable standing adduction", "Side-lying adduction", "Wide-stance squat for compound work"], placement="Late in leg sessions or in a hip-accessory circuit.", programming="2-4 sets of 12-20 reps, 45-90 seconds rest, 1-3 RIR.", progression="Increase reps and controlled range before adding one pin.", movement="thighs close inward"),
    guide(31, "Legs", "Standing Calf Raise Machine", [17], confidence="High", evidence="Shoulder pads, raised foot platform, plate horns and release mechanism identify a standing calf raise.", pattern="Ankle plantarflexion", level="Beginner to advanced", role="Calf isolation", purpose="Load standing ankle extension with the knees mostly straight.", primary="Gastrocnemius", secondary="Soleus and foot stabilizers", setup=["Load the horns evenly.", "Place the balls of the feet on the platform with heels free.", "Set shoulders under the pads, stand tall and learn the safety release before the set."], execution=["Lower heels into a controlled stretch.", "Rise as high as possible onto the balls of the feet.", "Pause without rolling ankles outward.", "Lower slowly and re-engage the safety before stepping out."], cues=["Straight up through big toes", "Soft knees", "Full controlled stretch"], mistakes=[("Bouncing", "Pause at bottom and reduce load."), ("Feet too far on platform", "Keep heels free."), ("Unlocking before balanced", "Set posture first.")], safety="Practice the release with an empty or light machine; secure it before leaving the shoulder pads.", alternatives=["Leg-press calf raise", "Smith calf raise", "Single-leg bodyweight calf raise"], placement="Late in leg training.", programming="3-5 sets of 8-20 reps, 60-120 seconds rest, 1-3 RIR.", progression="Add pauses and full range before increasing plates.", movement="heels rise and lower"),

    guide(32, "Core", "Adjustable Decline Core Benches", [6, 25], confidence="High", evidence="Both photographs show padded benches with leg rollers and adjustable decline positions.", pattern="Trunk flexion and supported floor exercise", level="Beginner to advanced", role="Core accessory station", purpose="Provide adjustable decline positions for crunches, sit-ups and selected dumbbell movements.", primary="Rectus abdominis for crunch variations", secondary="Obliques and hip flexors depending on exercise", setup=["Choose the shallowest useful decline and confirm the adjustment pin is locked.", "Secure lower legs under the rollers.", "Start with bodyweight and enough bench clearance."], execution=["For a crunch, exhale and bring ribs toward the pelvis.", "Pause before the lower back leaves the pad excessively.", "Lower one segment at a time.", "Use full sit-ups only if hip-flexor dominance and back comfort are acceptable."], cues=["Ribs toward pelvis", "Move slowly", "Use the bench angle, not momentum"], mistakes=[("Starting too steep", "Reduce the decline."), ("Pulling the neck", "Keep hands light or crossed."), ("Swinging", "Use a crunch range.")], safety="Confirm the bench pin and leg rollers are locked before lying back.", alternatives=["Exercise-mat dead bug", "Cable crunch", "Front plank"], placement="After compound training or in a brief core block.", programming="2-4 sets of 8-20 controlled reps, 45-90 seconds rest, stop before form changes.", progression="Add reps, then a slightly steeper angle, then a light load held at the chest.", movement="torso curls toward pelvis"),
    guide(33, "Core", "BOSU Balance Trainer", [48], confidence="High", evidence="Two half-dome BOSU-style balance trainers are visible.", pattern="Balance and low-load stability", level="Beginner with support to advanced under coaching", role="Balance or mobility accessory", purpose="Add an unstable surface to selected low-load balance, mobility and core drills.", primary="Exercise-dependent; ankle, hip and trunk stabilizers", secondary="Shoulders and legs depending on drill", setup=["Place the dome on a clean, non-slip floor with clear space around it.", "Begin beside a wall or fixed support.", "Use simple standing balance, step-ups or hands-elevated plank before complex drills."], execution=["Step onto the center deliberately.", "Maintain a tripod foot and soft knee.", "Use slow repetitions and step off before fatigue reduces control.", "Keep external load light unless supervised."], cues=["Own the surface", "Move slowly", "Use support early"], mistakes=[("Heavy squats on an unstable surface", "Use stable ground for strength loading."), ("Standing on damaged equipment", "Inspect first."), ("Progressing too quickly", "Increase time before complexity.")], safety="The dome increases fall risk; avoid near racks, weights or traffic lanes.", alternatives=["Single-leg floor balance", "Foam pad", "Controlled step-up"], placement="Warm-up, balance block or low-load accessory work.", programming="2-4 rounds of 20-45 second holds or 6-12 controlled reps.", progression="Reduce hand support, extend time or add slow head/arm movements before load.", movement="controlled balance"),
    guide(34, "Core", "Exercise Mats", [49], confidence="High", evidence="A rack of rolled exercise mats is clearly visible.", pattern="Floor support", level="All levels", role="Warm-up, core, mobility and recovery tool", purpose="Provide cushioning and grip for floor-based exercise without changing the exercise itself.", primary="Exercise-dependent", secondary="Exercise-dependent", setup=["Choose a clean, undamaged mat and place it on a flat floor.", "Allow enough space for arms and legs in every direction.", "Keep the mat away from cable paths, walkways and moving equipment."], execution=["Use for dead bugs, bird dogs, planks, glute bridges and mobility drills.", "Maintain the exercise-specific breathing and alignment.", "Wipe the mat after use.", "Roll it loosely and return it to the rack."], cues=["Clear space first", "Mat supports - it does not correct form", "Clean and re-rack"], mistakes=[("Using too much cushioning for standing lifts", "Use firm flooring for loaded balance."), ("Crowding traffic", "Move to an open area."), ("Leaving it unclean", "Follow facility hygiene rules.")], safety="A torn or slippery mat should be removed from use and reported.", alternatives=["Padded turf", "Flat bench for elevated drills", "Foam pad"], placement="Warm-up, core block, mobility or cooldown.", programming="Use according to the selected exercise; quality and breathing matter more than repetition speed.", progression="Progress the exercise, not the mat thickness.", movement="exercise-dependent"),

    guide(35, "Full Body", "Smith Machine", [18], confidence="High", evidence="A bar fixed to vertical guide rails with rotating hooks and plate storage is visible.", pattern="Fixed-path barbell training", level="Beginner to advanced with exercise-specific coaching", role="Multi-purpose compound station", purpose="Support squats, presses, lunges, rows, calf raises and hip thrusts on a guided bar path.", primary="Exercise-dependent", secondary="Exercise-dependent; stabilizer demand is lower than a free bar", setup=["Set safety stops before loading.", "Position body or bench so the fixed path fits the chosen exercise.", "Load both sides evenly and use collars if required by the facility."], execution=["Grip the bar and rotate it fully free of the hooks.", "Follow the exercise-specific path while respecting the fixed rail direction.", "Re-hook both sides together.", "Lower onto safeties only in an emergency or planned pin exercise."], cues=["Safeties first", "Match body to bar path", "Confirm both hooks"], mistakes=[("Assuming every free-bar setup transfers directly", "Adjust stance or bench to the fixed rails."), ("No safeties", "Set them before loading."), ("One hook caught", "Visually confirm both sides.")], safety="Never rely only on wrist rotation; safeties must be positioned for presses and squats.", alternatives=["Power rack", "Leg press", "Machine presses"], placement="As the main compound station or a controlled accessory.", programming="Depends on exercise; most compounds use 3-5 sets of 5-12 reps with 1-3 RIR.", progression="Progress the selected exercise while keeping safety-stop and body positions recorded.", movement="fixed vertical bar path"),
    guide(36, "Full Body", "Power Rack / Half Rack", [29, 30], confidence="High", evidence="Uprights, J-hooks, pull-up grips, plate storage and alternate front/side views identify the same rack system.", pattern="Free-barbell training", level="Intermediate skill demand", role="Multi-purpose strength station", purpose="Support squats, bench presses, overhead presses, rack pulls and pull-ups with adjustable safeties.", primary="Exercise-dependent", secondary="Core, grip and stabilizers", setup=["Set J-hooks just below the finished unrack position.", "Set safety arms just below the lowest controlled position.", "Center the bar, load evenly and apply collars as appropriate."], execution=["Unrack with a stable brace and minimal extra steps.", "Perform the exercise within the safety zone.", "Re-rack by contacting both uprights and lowering onto both hooks.", "Unload plates evenly, alternating sides when heavy."], cues=["Hooks at the right height", "Safeties catch the miss", "Re-rack both sides"], mistakes=[("Safeties too low", "Test with an empty bar."), ("Walking backward to re-rack", "Contact uprights first."), ("Stripping one side completely", "Alternate plates to prevent tipping.")], safety="Use a spotter for challenging lifts even when safeties are present; inspect J-hooks before loading.", alternatives=["Smith machine", "Machine compound lifts", "Dumbbell training"], placement="Primary station for free-barbell workouts.", programming="Exercise-specific; use conservative loads until technique is reliable.", progression="Increase weight in small steps while preserving setup measurements and bar path.", movement="exercise-dependent free bar"),
    guide(37, "Full Body", "Matrix Functional Trainer", [19, 20], confidence="High", evidence="Matrix dual adjustable pulleys, two weight stacks and mirrored alternate views are visible.", pattern="Multi-planar cable training", level="Beginner to advanced", role="Multi-purpose station", purpose="Provide adjustable cable resistance for presses, rows, curls, extensions, raises and trunk work.", primary="Exercise-dependent", secondary="Core and stabilizers", setup=["Move both pulley carriages to the chosen height and confirm their locks.", "Attach handles with fully closed carabiners.", "Select equal stack loads for bilateral exercises unless intentionally asymmetric."], execution=["Step far enough away to maintain cable tension.", "Use a stable stance and move through the intended joint action.", "Control the cable back to the start.", "Return carriages and attachments after the set."], cues=["Cable follows the limb", "Stand beyond slack", "Quiet stacks"], mistakes=[("Unlocked carriage", "Test the lock before loading."), ("Standing in another cable's path", "Check both sides."), ("Letting handles snap back", "Walk them to the start.")], safety="Inspect carabiners and never cross an occupied cable lane.", alternatives=["FreeMotion dual cable cross", "Selectorized machines", "Resistance bands"], placement="Main movement, accessory circuit, warm-up or rehabilitation-style low-load work.", programming="2-4 sets of 8-20 reps per exercise; rest 45-120 seconds based on complexity.", progression="Add reps, cable travel or one stack increment while preserving position.", movement="adjustable cable path"),
    guide(38, "Full Body", "FreeMotion Genesis Dual Cable Cross", [24], confidence="High", evidence="The front panel reads FreeMotion Genesis Dual Cable Cross and articulating cable arms are visible.", pattern="Multi-planar cable training", level="Beginner to advanced", role="Multi-purpose functional station", purpose="Use independently adjustable arms and stacks for nearly any cable movement direction.", primary="Exercise-dependent", secondary="Core and stabilizers", setup=["Position each articulating arm and confirm every lock.", "Attach handles to the correct cables.", "Select loads and stand where both cables remain tensioned."], execution=["Set a stable stance.", "Perform the chosen press, row, chop, curl, extension or lower-body movement.", "Keep cable angles aligned with the intended movement.", "Guide handles back before changing arm positions."], cues=["Lock every joint", "Cables point through the movement", "Control both handles"], mistakes=[("Asymmetric arm positions for bilateral work", "Match the settings."), ("Too much load for an unstable stance", "Reduce load or widen stance."), ("Changing arms under tension", "Return stacks first.")], safety="Never release an arm-position lock while a cable is loaded or a person is in the cable path.", alternatives=["Matrix functional trainer", "Single cable column", "Dumbbells"], placement="Any point in the session depending on the exercise.", programming="2-4 sets of 8-20 reps; compound cable movements may use longer rests.", progression="Progress load, repetitions or stance complexity one variable at a time.", movement="articulating cable path"),
    guide(39, "Full Body", "Dumbbells", [22], confidence="High", evidence="A two-tier rack of light through heavy hex dumbbells is clearly visible.", pattern="Free-weight multi-planar training", level="All levels", role="Compound and isolation tool", purpose="Allow independent-arm loading for presses, rows, carries, squats, lunges, hinges, curls and raises.", primary="Exercise-dependent", secondary="Grip and stabilizers", setup=["Read both dumbbell labels and choose a matched pair.", "Clear floor space and position a bench if needed.", "For heavy dumbbells, use a safe two-hand lift and controlled knee assist into position."], execution=["Keep wrists stacked and each weight balanced.", "Follow the selected exercise's stable path.", "End the set before control is lost.", "Place weights down deliberately and re-rack in labeled order."], cues=["Own each weight independently", "No dropped dumbbells", "Match the pair"], mistakes=[("Using mismatched weights", "Check both labels."), ("Throwing dumbbells", "Use a lighter load and controlled exit."), ("Blocking the rack", "Step away after selecting weights.")], safety="Do not leave dumbbells where they create a trip hazard; inspect loose heads or damaged handles.", alternatives=["Barbells", "Cable training", "Selectorized machines"], placement="Any point depending on exercise complexity.", programming="Compounds: 3-5 sets of 6-15; isolation: 2-4 sets of 10-20.", progression="Use double progression: reach the top rep target, then move to the next pair.", movement="exercise-dependent free weight"),
    guide(40, "Full Body", "Light Dumbbells and Small Tools", [47], confidence="High", evidence="Color-coded light dumbbells, medicine balls and organized small-tool storage are visible.", pattern="Low-load resistance and conditioning", level="All levels", role="Warm-up, accessory and class tools", purpose="Support low-load raises, rotator-cuff drills, carries, mobility and conditioning circuits.", primary="Exercise-dependent", secondary="Grip and stabilizers", setup=["Select the lightest tool that allows the intended control.", "Clear enough space for the whole movement.", "Check balls, handles and coatings for damage."], execution=["Use deliberate repetitions rather than momentum.", "Keep range matched to the target joint and exercise.", "Stop before small stabilizers lose position.", "Return every item to its labeled compartment."], cues=["Light means precise", "Smooth arcs", "Re-rack by color and weight"], mistakes=[("Using speed to create difficulty", "Slow the tempo."), ("Choosing load by color only", "Read the label."), ("Leaving tools on turf", "Return them promptly.")], safety="Small equipment becomes a trip hazard quickly; keep the training area organized.", alternatives=["Cable accessories", "Resistance bands", "Bodyweight drills"], placement="Warm-up, accessory circuit or cooldown.", programming="2-4 sets of 10-25 controlled reps or timed circuits of 20-45 seconds.", progression="Add control, range or time before choosing a heavier tool.", movement="exercise-dependent low load"),
    guide(41, "Full Body", "Weight Plates and Loading Practices", [7], confidence="Medium", evidence="The photograph is a plate-loaded-machine area overview with multiple plate horns; it is intentionally treated as an overview rather than a single-machine identification.", pattern="Loading and facility safety", level="All levels", role="Equipment-use skill", purpose="Teach balanced loading, plate handling, collars and safe unloading across plate-loaded machines.", primary="Not an exercise", secondary="Grip and lifting mechanics during handling", setup=["Read plate labels and inspect the storage horn.", "Carry one manageable plate close to the body with both hands when needed.", "Load the same amount on both working horns unless a unilateral design specifically permits otherwise."], execution=["Slide plates fully onto the horn.", "Add collars where the machine or facility requires them.", "After training, unload large outer plates first and alternate sides.", "Return every plate to its matching storage location."], cues=["Even sides", "Plate fully seated", "Alternate while unloading"], mistakes=[("Treating the overview as one machine", "Use its visible plate stations only for loading instruction."), ("Stripping one heavy side", "Alternate sides."), ("Mixing plate sizes on storage", "Use labeled horns.")], safety="Never put fingers between stacked plates or stand under a loaded moving lever.", alternatives=["Selectorized stack", "Fixed dumbbells", "Resistance bands"], placement="Before and after every plate-loaded exercise.", programming="Not programmed as an exercise; use sound handling on every set.", progression="Increase training load only after setup, balance and unloading habits are automatic.", movement="safe plate handling"),
    guide(42, "Full Body", "Adjustable Utility Benches", [31, 32, 33], confidence="High", evidence="Flat, incline and decline bench configurations are shown within dedicated barbell stations; the guide focuses on bench angle and locking checks.", pattern="Support surface", level="All levels", role="Multi-purpose support equipment", purpose="Provide stable flat, inclined or declined body positions for dumbbell, barbell and bodyweight exercises.", primary="Exercise-dependent", secondary="Exercise-dependent", setup=["Inspect the frame, pad and adjustment ladder or pin.", "Lock the selected angle before sitting or adding weight.", "Center the bench relative to racks, cables or open floor space."], execution=["Sit before bringing heavy dumbbells into position.", "Maintain full pad contact appropriate to the exercise.", "Exit with weights controlled and feet clear of the frame.", "Return adjustable benches to the facility's standard position."], cues=["Lock before load", "Center the bench", "Clear feet and fingers"], mistakes=[("Half-engaged ladder", "Lift and reseat it fully."), ("Bench blocking rack access", "Move only to designated areas."), ("Standing on upholstery", "Use a stable platform instead.")], safety="Never adjust the backrest while someone is lying on the bench or holding weight.", alternatives=["Fixed bench station", "Floor exercise", "Machine back pad"], placement="As required by the chosen exercise.", programming="The bench supports the exercise; follow that exercise's prescription.", progression="Progress the exercise while preserving identical bench angle and position.", movement="support position"),

    guide(43, "Cardio", "Technogym Excite Top Upper-Body Ergometer", [50], confidence="High", evidence="Technogym branding and the seated hand-crank unit in the foreground match the Excite Top upper-body ergometer; it is not a StepMill.", pattern="Upper-body cyclical cardio", level="Beginner to advanced", role="Warm-up, conditioning or accessible cardio", purpose="Create cardiovascular work through adjustable arm cranking while seated or positioned at the unit.", primary="Cardiovascular system; shoulders and arms perform the cyclical work", secondary="Upper back, chest, forearms and trunk stabilizers", setup=["Adjust seat distance and crank height if the unit permits.", "Begin with low resistance and confirm both handles rotate freely.", "Sit tall with shoulders relaxed and wrists neutral."], execution=["Crank in a smooth rhythm using both arms.", "Keep the torso stable and avoid shrugging.", "Increase resistance or cadence gradually.", "Cool down at low resistance before stopping."], cues=["Smooth circles", "Shoulders down", "Build intensity gradually"], mistakes=[("Calling it a StepMill", "The photographed unit is an upper-body ergometer."), ("Gripping too hard", "Relax the hands."), ("Starting at high resistance", "Warm up first.")], safety="Stop for chest pain, faintness or unusual shortness of breath and follow the facility's emergency process.", alternatives=["Recumbent bike", "Treadmill walking", "Light cable circuit"], placement="Warm-up, standalone cardio or lower-body recovery day.", programming="Start with 5-15 minutes easy; progress toward 20-30 minutes or short intervals as tolerated.", progression="Increase duration first, then cadence or resistance, one variable at a time.", movement="arm-crank rotation"),
    guide(44, "Cardio", "Recumbent Exercise Bikes", [50], confidence="High", evidence="A row of back-supported recumbent bikes is visible behind the Technogym upper-body ergometer in the cardio-area photograph.", pattern="Lower-body cyclical cardio", level="Beginner to advanced", role="Warm-up or conditioning", purpose="Provide low-impact seated cycling with a supportive backrest.", primary="Cardiovascular system and quadriceps", secondary="Glutes, hamstrings and calves", setup=["Choose an available bike in the background row and adjust the seat.", "At full pedal extension, keep a small bend in the knee.", "Set low resistance before beginning."], execution=["Pedal smoothly without rocking the hips.", "Maintain a cadence that matches the planned intensity.", "Increase resistance gradually.", "Cool down before bringing pedals to a stop."], cues=["Knee stays softly bent", "Hips stay on seat", "Smooth cadence"], mistakes=[("Seat too close", "Move back until the knee has a small bend."), ("High resistance immediately", "Warm up first."), ("Pushing only with toes", "Keep foot pressure balanced.")], safety="Use the seat lock and stop if knee or hip pain appears.", alternatives=["Treadmill walking", "Upper-body ergometer", "Easy outdoor cycling"], placement="Warm-up, recovery cardio or standalone aerobic session.", programming="10-40 minutes at conversational effort, or short intervals after a gradual warm-up.", progression="Add 5 minutes before increasing resistance or interval density.", movement="pedal rotation"),
    guide(45, "Cardio", "Treadmills", [51], confidence="High", evidence="A row of Technogym treadmills with belts, consoles and emergency-stop controls is visible.", pattern="Walking or running", level="Beginner to advanced", role="Warm-up, aerobic training or intervals", purpose="Support controlled indoor walking, incline walking and running.", primary="Cardiovascular system; lower body performs locomotion", secondary="Quadriceps, hamstrings, glutes, calves and trunk", setup=["Stand on the side rails before starting the belt.", "Attach the safety clip if provided and select a very low starting speed.", "Step onto the moving belt only after it is stable."], execution=["Walk near the center with natural arm swing.", "Increase speed or incline gradually.", "Use rails briefly for balance, not to support bodyweight continuously.", "Reduce speed for a cooldown before stepping to the side rails."], cues=["Tall posture", "Quiet steps", "Change one setting at a time"], mistakes=[("Starting on the belt", "Begin from the rails."), ("Holding rails while using steep incline", "Lower intensity to one you can support."), ("Looking at feet continuously", "Look ahead once stable.")], safety="Know the emergency stop; stop for chest pain, dizziness or unusual shortness of breath.", alternatives=["Recumbent bike", "Upper-body ergometer", "Outdoor walking"], placement="Warm-up, standalone cardio, intervals or cooldown.", programming="Begin with 10-20 minutes easy; build toward public-health aerobic targets over time.", progression="Increase weekly minutes before speed or incline; avoid simultaneous large jumps.", movement="moving walking belt"),
]


SOURCES = [
    ("American College of Sports Medicine. Resistance Training Guidelines Update (2026).", "https://acsm.org/resistance-training-guidelines-update-2026/"),
    ("American College of Sports Medicine. Resistance Training Position Stand Infographic (2026).", "https://acsm.org/wp-content/uploads/2026/03/Resistance-Training-Position-Stand-infographic.pdf"),
    ("U.S. Department of Health and Human Services. Physical Activity Guidelines for Americans, 2nd edition.", "https://odphp.health.gov/our-work/nutrition-physical-activity/physical-activity-guidelines/current-guidelines"),
    ("U.S. HHS. Top 10 Things to Know About the Physical Activity Guidelines.", "https://odphp.health.gov/our-work/nutrition-physical-activity/physical-activity-guidelines/current-guidelines/top-10-things-know"),
    ("USDA and HHS. Dietary Guidelines for Americans, 2025-2030 portal and history.", "https://www.dietaryguidelines.gov/history"),
    ("Morton RW et al. Protein supplementation and resistance training adaptations. Br J Sports Med. 2018.", "https://pubmed.ncbi.nlm.nih.gov/28698222/"),
    ("Jager R et al. ISSN Position Stand: Protein and Exercise. JISSN. 2017.", "https://doi.org/10.1186/s12970-017-0177-8"),
    ("Matrix Fitness. Engineering a Smarter, More Stable Take on Glute Training.", "https://us.matrixfitness.com/eng/innovations/glute-trainer"),
    ("Booty Builder. Platinum Hip Thrust Machine product information.", "https://bootybuilder.com/product/booty-builder-platinum/"),
    ("Life Fitness Support. Insignia Series Strength Owner's Manual index.", "https://support.lifefitness.com/hc/en-us/articles/360043013933-Life-Fitness-Insignia-Series-Strength-Owner-s-Manual"),
    ("Life Fitness Support. Circuit Series Machine Use Instructions.", "https://support.lifefitness.com/hc/en-us/articles/360036936034-Circuit-Series-Machine-Use-Instructions"),
    ("FreeMotion Fitness. Company history and Genesis DS dual-cable system.", "https://www.freemotionfitness.com/en-gb/about-us/"),
    ("Technogym Annual Report product range: Excite Top upper-body ergometer.", "https://corporate.technogym.com/~/media/Files/T/Technogym-Corporate/reports-and-presentation/2024/annual-report-2024-en.pdf"),
]


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # Treat the first row as the table's accessible header. For one-row visual
    # callouts this gives assistive technology a labelled region instead of an
    # anonymous layout table; for data tables it also enables repeated headers.
    set_repeat_table_header(table.rows[0])


def set_table_borders(table, color: str = "C7D2CC", size: int = 4):
    """Add a light, printable grid without relying on a Word table style."""
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_row_height(row, height_dxa: int = 500):
    tr_pr = row._tr.get_or_add_trPr()
    height = OxmlElement("w:trHeight")
    height.set(qn("w:val"), str(height_dxa))
    height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(height)


def add_bookmark(paragraph, name: str, bookmark_id: int):
    safe = re.sub(r"[^A-Za-z0-9_]", "_", name)[:35]
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), safe)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)
    return safe


def add_internal_link(paragraph, text: str, anchor: str, color: str = "2E6FA7"):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_external_link(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "2E6FA7")
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def set_run(run, size=10, bold=False, color=INK, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)


def set_alt_text(shape, text: str):
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", text)
    doc_pr.set("title", text[:100])


def add_field(paragraph, field_code: str):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2, placeholder, fld_char3])
    return run


def rotate_and_fit(src: Path, size=(1600, 980), crop=None) -> Image.Image:
    im = ImageOps.exif_transpose(Image.open(src)).convert("RGB")
    if crop:
        x0, y0, x1, y1 = crop
        im = im.crop((int(im.width*x0), int(im.height*y0), int(im.width*x1), int(im.height*y1)))
    bg = Image.new("RGB", size, "#10171D")
    fitted = ImageOps.contain(im, (size[0]-220, size[1]-40), Image.Resampling.LANCZOS)
    x = (size[0]-fitted.width)//2
    y = (size[1]-fitted.height)//2
    bg.paste(fitted, (x, y))
    return bg


def font(path: str, size: int):
    try:
        return ImageFont.truetype(path, size)
    except Exception:
        return ImageFont.load_default()


FONT_BOLD = r"C:\Windows\Fonts\arialbd.ttf"
FONT_REG = r"C:\Windows\Fonts\arial.ttf"


def draw_arrow(draw: ImageDraw.ImageDraw, start, end, color, width=8):
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1]-start[1], end[0]-start[0])
    length = 24
    for delta in (2.55, -2.55):
        p = (end[0] + length*math.cos(angle+delta), end[1] + length*math.sin(angle+delta))
        draw.line([end, p], fill=color, width=width)


def annotated_asset(g: dict, photo_num: int, suffix: str = "") -> Path:
    src = SOURCE_DIR / photo_name(photo_num)
    out = ASSET_DIR / f"guide_{g['no']:02d}_{photo_num}{suffix}_annotated.jpg"
    crop = None
    if g["no"] == 44:
        crop = (0.18, 0.18, 0.90, 0.78)
    im = rotate_and_fit(src, crop=crop)
    draw = ImageDraw.Draw(im, "RGBA")
    color = "#" + CATEGORY_COLORS[g["category"]]
    positions = [
        ((35, 90), (550, 330)),
        ((1030, 90), (920, 430)),
        ((35, 820), (560, 700)),
        ((1030, 820), (1000, 680)),
    ]
    labels = g["annotations"][:4]
    for idx, label in enumerate(labels):
        box, target = positions[idx]
        x, y = box
        tw = min(480, max(240, 22*len(label)))
        draw.rounded_rectangle((x, y, x+tw, y+58), radius=14, fill=(16, 23, 29, 225), outline=color, width=4)
        draw.text((x+18, y+14), label.upper(), font=font(FONT_BOLD, 24), fill="white")
        start = (x+tw if x < 800 else x, y+29)
        draw_arrow(draw, start, target, color, 7)
    draw.rounded_rectangle((540, 435, 1060, 505), radius=14, fill=(16, 23, 29, 215), outline=color, width=4)
    draw.text((565, 454), f"MOVEMENT: {g['movement'].upper()}", font=font(FONT_BOLD, 25), fill="white")
    draw_arrow(draw, (790, 520), (790, 635), color, 10)
    im.save(out, quality=88, optimize=True)
    return out


def plain_asset(photo_num: int, suffix: str = "") -> Path:
    out = ASSET_DIR / f"photo_{photo_num}{suffix}.jpg"
    if out.exists():
        return out
    crop = (0.18, 0.18, 0.90, 0.78) if suffix == "_bikes" else None
    rotate_and_fit(SOURCE_DIR / photo_name(photo_num), crop=crop).save(out, quality=88, optimize=True)
    return out


def cover_asset() -> Path:
    out = ASSET_DIR / "cover_collage.jpg"
    panels = [rotate_and_fit(SOURCE_DIR / photo_name(n), size=(800, 600)) for n in (0, 11, 24, 51)]
    canvas = Image.new("RGB", (1600, 1200), "#10171D")
    for i, panel in enumerate(panels):
        canvas.paste(panel, ((i%2)*800, (i//2)*600))
    overlay = Image.new("RGBA", canvas.size, (9, 18, 25, 85))
    canvas = Image.alpha_composite(canvas.convert("RGBA"), overlay).convert("RGB")
    canvas.save(out, quality=88, optimize=True)
    return out


def configure_styles(doc: Document):
    sec = doc.sections[0]
    sec.page_width = PAGE_W
    sec.page_height = PAGE_H
    sec.top_margin = MARGIN
    sec.bottom_margin = MARGIN
    sec.left_margin = MARGIN
    sec.right_margin = MARGIN
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in (
        ("Title", 30, INK, 0, 8),
        ("Subtitle", 14, MUTED, 0, 8),
        ("Heading 1", 16, "2E6FA7", 12, 7),
        ("Heading 2", 13, "2E6FA7", 9, 5),
        ("Heading 3", 11.5, "1F4D78", 7, 4),
    ):
        s = styles[name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        s.font.size = Pt(size)
        s.font.color.rgb = rgb(color)
        s.font.bold = name != "Subtitle"
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        s = styles[style_name]
        s.font.name = "Calibri"
        s.font.size = Pt(9.5)
        s.paragraph_format.left_indent = Inches(0.375)
        s.paragraph_format.first_line_indent = Inches(-0.188)
        s.paragraph_format.space_after = Pt(3)
        s.paragraph_format.line_spacing = 1.08

    if "Guide Body" not in styles:
        s = styles.add_style("Guide Body", WD_STYLE_TYPE.PARAGRAPH)
    else:
        s = styles["Guide Body"]
    s.font.name = "Calibri"
    s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    s.font.size = Pt(9.5)
    s.font.color.rgb = rgb(INK)
    s.paragraph_format.space_after = Pt(3)
    s.paragraph_format.line_spacing = 1.08

    if "Small Caption" not in styles:
        s = styles.add_style("Small Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        s = styles["Small Caption"]
    s.font.name = "Calibri"
    s.font.size = Pt(8)
    s.font.color.rgb = rgb(MUTED)
    s.font.italic = True
    s.paragraph_format.space_after = Pt(4)


def configure_header_footer(doc: Document):
    sec = doc.sections[0]
    header = sec.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("THE COMPLETE GYM EQUIPMENT HANDBOOK")
    set_run(run, size=8, bold=True, color=MUTED)
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), RULE)
    pbdr.append(bottom)
    p_pr.append(pbdr)

    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("PAGE ")
    set_run(r, size=8, bold=True, color=MUTED)
    r = add_field(p, "PAGE")
    set_run(r, size=8, bold=True, color=MUTED)
    r = p.add_run(" OF ")
    set_run(r, size=8, bold=True, color=MUTED)
    r = add_field(p, "NUMPAGES")
    set_run(r, size=8, bold=True, color=MUTED)


PAGE_COUNT = 0
BOOKMARK_ID = 1
BOOKMARKS: dict[str, str] = {}
PAGE_NUMBERS: dict[str, int] = {}


def start_page(doc: Document, title: str, *, kicker: str = "", color: str = "2E6FA7", bookmark: str | None = None, subtitle: str = ""):
    global PAGE_COUNT, BOOKMARK_ID
    if PAGE_COUNT:
        doc.add_page_break()
    PAGE_COUNT += 1
    if kicker:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.keep_with_next = True
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), color)
        p_pr.append(shd)
        run = p.add_run("  " + kicker.upper() + "  ")
        set_run(run, size=8, bold=True, color=WHITE)
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    set_run(r, size=16, bold=True, color=INK)
    if bookmark:
        safe = add_bookmark(p, bookmark, BOOKMARK_ID)
        BOOKMARK_ID += 1
        BOOKMARKS[bookmark] = safe
        PAGE_NUMBERS[bookmark] = PAGE_COUNT
    if subtitle:
        p = doc.add_paragraph(style="Subtitle")
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(subtitle)
        set_run(r, size=11, color=MUTED)
    return p


def start_section_on_current_page(doc: Document, title: str, *, kicker: str, color: str, bookmark: str):
    """Add a bookmarked major section without forcing another page break."""
    global BOOKMARK_ID
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(8)
    spacer.paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    p_pr.append(shd)
    run = p.add_run("  " + kicker.upper() + "  ")
    set_run(run, size=8, bold=True, color=WHITE)
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run(title), size=16, bold=True, color=INK)
    safe = add_bookmark(p, bookmark, BOOKMARK_ID)
    BOOKMARK_ID += 1
    BOOKMARKS[bookmark] = safe
    PAGE_NUMBERS[bookmark] = PAGE_COUNT


def add_label_value(doc, label: str, value: str, color: str = "2E6FA7"):
    p = doc.add_paragraph(style="Guide Body")
    p.paragraph_format.keep_together = True
    r = p.add_run(label.upper() + "  ")
    set_run(r, size=8.5, bold=True, color=color)
    r = p.add_run(value)
    set_run(r, size=9.5, color=INK)
    return p


def add_bullets(doc, items: list[str], numbered=False, size=9.5):
    for index, item in enumerate(items, 1):
        if numbered:
            # Manual labels make every independent procedure restart at 1 in
            # Word and PDF, avoiding document-wide automatic-list continuation.
            p = doc.add_paragraph(style="Guide Body")
            p.paragraph_format.left_indent = Inches(0.28)
            p.paragraph_format.first_line_indent = Inches(-0.22)
            r = p.add_run(f"{index}. ")
            set_run(r, size=size, bold=True, color=INK)
        else:
            p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run(r, size=size, color=INK)


def add_callout(doc, heading: str, text: str, color: str, fill: str = LIGHT):
    body_children = list(doc._body._body)
    previous = next((element for element in reversed(body_children) if element.tag != qn("w:sectPr")), None)
    if previous is not None and previous.tag == qn("w:tbl"):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.line_spacing = Pt(1)
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_W_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(heading.upper() + "  ")
    set_run(r, size=9, bold=True, color=color)
    r = p.add_run(text)
    set_run(r, size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_image(doc, path: Path, width=6.25, alt=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    set_alt_text(shape, alt or path.stem)
    return shape


def add_two_images(doc, paths: list[Path], alts: list[str]):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [4680, 4680], indent_dxa=120)
    for i, path in enumerate(paths[:2]):
        cell = table.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape = p.add_run().add_picture(str(path), width=Inches(3.05))
        set_alt_text(shape, alts[i])
    return table


def make_cover(doc: Document):
    global PAGE_COUNT
    PAGE_COUNT = 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("THE COMPLETE")
    set_run(r, size=12, bold=True, color=GOLD)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("GYM EQUIPMENT")
    set_run(r, size=31, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("HANDBOOK")
    set_run(r, size=31, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("A photo-mapped reference for safe setup, sound technique and practical programming")
    set_run(r, size=11, color=MUTED)
    add_image(doc, cover_asset(), width=6.25, alt="Collage of photographed gym equipment from the handbook")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("45 GUIDES  |  51 SOURCE PHOTOGRAPHS  |  BEGINNER TO INTERMEDIATE")
    set_run(r, size=9, bold=True, color=GREEN)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Revised edition - 2026")
    set_run(r, size=9, color=MUTED)


def add_front_matter(doc: Document):
    start_page(doc, "Publication Note and Safety Disclaimer", kicker="Read first", color=RED, bookmark="disclaimer", subtitle="How to use this manual responsibly")
    add_callout(doc, "Not medical care", "This handbook is educational. It does not diagnose injury, prescribe rehabilitation or replace advice from a physician, physical therapist or qualified coach.", RED, "FBEDED")
    add_bullets(doc, [
        "Stop immediately for sharp pain, chest pain, faintness, new numbness, loss of coordination or unusual shortness of breath.",
        "Ask gym staff to demonstrate unfamiliar safety stops, seat mechanisms, counterweights and emergency controls.",
        "The photographed machines may be serviced, moved or replaced. The placard on the current machine always controls when it conflicts with this handbook.",
        "Begin with conservative resistance. A successful first session should feel repeatable, not maximal.",
        "People who are pregnant, returning after surgery, managing chronic disease or experiencing persistent pain should obtain individualized guidance.",
    ])
    add_callout(doc, "Emergency rule", "Know the location of staff, emergency stop controls and the facility's emergency response procedure before training alone.", RED, "FBEDED")
    add_label_value(doc, "Edition scope", "Healthy beginner-to-intermediate adults; strength, conditioning and general fitness.")
    add_label_value(doc, "Image policy", "All 51 source filenames are mapped in the photo index. Alternate angles are grouped, not silently omitted.")

    start_page(doc, "How to Use This Handbook", kicker="Navigation", color="2E6FA7", bookmark="how_to_use", subtitle="A three-pass method for every unfamiliar station")
    add_bullets(doc, [
        "PASS 1 - Identify: match the machine to the guide title and source photo caption. Check the confidence note.",
        "PASS 2 - Set up: find the seat, pads, start-angle lever, selector pin, safety stops and attachment points before adding effort.",
        "PASS 3 - Perform: use the numbered sequence, form cues and stop conditions. Record successful settings for next time.",
    ], numbered=True)
    add_callout(doc, "RIR", "Repetitions in reserve estimates how many technically sound repetitions were still possible. Most work in this handbook ends with about 1-3 RIR.", "2E6FA7")
    add_label_value(doc, "Guide page 1", "Identification, evidence, annotated photograph and machine role.")
    add_label_value(doc, "Guide page 2", "Adjustments, setup, execution, muscles, breathing and form cues.")
    add_label_value(doc, "Guide page 3", "Programming, progression, mistakes, alternatives and workout placement.")
    add_callout(doc, "Photo captions", "Captions preserve the exact source filename so every photograph remains traceable.", GREEN, "EDF7F1")

    start_page(doc, "Visual Key", kicker="Symbols and colors", color="7356A5", bookmark="visual_key", subtitle="What the repeated treatments mean")
    for cat, color in CATEGORY_COLORS.items():
        add_callout(doc, cat, {
            "Push": "Chest, shoulder and triceps pressing.",
            "Pull": "Back, rear-shoulder and elbow-flexion training.",
            "Legs": "Knee, hip, calf and inner-thigh training.",
            "Core": "Trunk, balance and floor-support equipment.",
            "Full Body": "Racks, cables, free weights and general tools.",
            "Cardio": "Walking, cycling and upper-body aerobic equipment.",
        }[cat], color, "F7F9FA")
    add_label_value(doc, "Red warning", "A condition that should stop the set or requires additional care.", RED)
    add_label_value(doc, "Green check", "A repeatable setup or progression milestone.", GREEN)

    start_page(doc, "Identification and Confidence Method", kicker="Evidence", color=GOLD, bookmark="id_method", subtitle="Names are based on visible evidence, not guesswork")
    add_bullets(doc, [
        "HIGH: the exercise placard, brand/model cue and machine geometry agree.",
        "MEDIUM-HIGH: geometry and adjustment layout are clear, but the exact model text is not readable.",
        "MEDIUM: the photograph is an area overview or the exact product variant cannot be confirmed.",
        "No machine receives a brand/model claim solely because another nearby machine has that label.",
    ])
    add_callout(doc, "Corrected mappings", "Image 3 is the plate-loaded pulldown; image 4 is a press; image 9 is a Matrix glute trainer; image 12 is a lying leg curl; images 44-45 both show hip adduction; image 50 shows a Technogym upper-body ergometer and recumbent bikes, not a StepMill.", RED, "FBEDED")
    add_label_value(doc, "Primary evidence", "Visible placards, manufacturer marks, user contact points, pivots, cable/lever direction, load system and safety hardware.")
    add_label_value(doc, "Conservative rule", "When a model is uncertain, the guide uses a generic equipment name and tells the reader to follow the current placard.")

    start_page(doc, "Photo-Mapping Overview", kicker="51 of 51", color=GREEN, bookmark="photo_overview", subtitle="Every source photograph is included or explicitly cross-referenced")
    counts = {c: 0 for c in CATEGORY_COLORS}
    for g in GUIDES:
        counts[g["category"]] += 1
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Section", "Guides", "Photo treatment"
    for cell in hdr:
        set_cell_shading(cell, "E8EEF5")
    for cat in CATEGORY_COLORS:
        cells = table.add_row().cells
        cells[0].text = cat
        cells[1].text = str(counts[cat])
        cells[2].text = "Primary image, alternate angle or indexed cross-reference"
    set_table_geometry(table, [2100, 1200, 6060])
    set_repeat_table_header(table.rows[0])
    add_callout(doc, "Coverage rule", "A photograph can support more than one guide when it contains multiple clearly visible stations. Its original filename appears in each relevant caption/index entry.", GREEN, "EDF7F1")
    add_label_value(doc, "Distinct guides", "45")
    add_label_value(doc, "Source photographs", "51")
    add_label_value(doc, "Unsupported StepMill claims", "0")


PRINCIPLES = [
    ("Training Readiness and Stop Signs", "readiness", RED, ["Use a readiness check before loading", "Stop conditions are different from normal effort", "Pain is information, not a challenge"], "Normal training effort includes working-muscle fatigue, faster breathing and a controlled rise in heart rate. Stop for sharp or escalating joint pain, chest pressure, faintness, new numbness, sudden weakness, loss of balance or symptoms that feel medically unusual. A conservative decision protects the next session."),
    ("Emergency Controls and Staff Support", "emergency", RED, ["Locate staff and exits", "Learn treadmill emergency stops", "Report damaged pads, cables or pins"], "Before a first session, identify the staffed desk, emergency phone or call point, automated external defibrillator location if posted, and the stop controls on cardio equipment. Never improvise a repair. Tag or report a damaged machine and choose another exercise."),
    ("Universal Machine Setup", "machine_setup", "2E6FA7", ["Joint near pivot", "Pad contacts the intended body segment", "Handles permit neutral wrists", "Full range remains pain-free"], "Most selectorized machines work best when the user's primary joint is close to the machine pivot. Set the seat first, then pads and start range, then resistance. Test one unloaded or very light repetition before the work set."),
    ("Loading, Collars, Safeties and Pins", "loading", GOLD, ["Selector pin fully inserted", "Plate horns loaded evenly", "Safeties tested with an empty bar", "Attachments and carabiners closed"], "Resistance is only useful when it stays attached and controllable. Record seat and safety settings as carefully as the weight. On racks, test safety height before loading. On plate-loaded machines, alternate sides while unloading heavy plates."),
    ("Posture, Bracing and Breathing", "breathing", "2E6FA7", ["Stack ribs over pelvis", "Breathe into the trunk", "Exhale through the hard portion", "Do not hold breath longer than needed"], "For most moderate sets, inhale during the easier or lowering phase and exhale through the hardest portion. A brief brace can improve trunk stiffness, but prolonged breath-holding is not necessary for most beginners and may be inappropriate for some health conditions."),
    ("Tempo, Range of Motion and RIR", "tempo", "7356A5", ["Use a controlled lowering phase", "Pause instead of bouncing", "Use the largest comfortable range you can control", "Finish most sets with 1-3 RIR"], "Tempo is a control tool, not a magic formula. A two-to-three second eccentric makes poor positions easier to notice. Range should be individualized: more is useful only while joint position, pad contact and control remain sound."),
    ("Movement Patterns", "patterns", "C67A24", ["Horizontal push and pull", "Vertical push and pull", "Knee-dominant and hip-dominant lower body", "Carry, rotation control and cyclical cardio"], "A balanced program trains patterns rather than chasing every machine. Use the handbook categories to select one or two main patterns per session, then add targeted accessories for muscles or skills that need more work."),
    ("Practical Muscle Map", "muscle_map", "7356A5", ["Chest, shoulders and triceps press", "Lats, upper back and biceps pull", "Quadriceps extend the knee", "Glutes extend the hip", "Hamstrings flex the knee and assist hip extension"], "Machine placards simplify anatomy, but no exercise isolates a single structure perfectly. Primary movers create most of the force; secondary muscles assist or stabilize. Choose exercises by movement and tolerance, not by an exaggerated promise to isolate one small region."),
    ("General Warm-Up", "warmup", RED, ["3-8 minutes easy cyclical movement", "Move major joints through comfortable ranges", "Practice the first exercise with light resistance", "Increase readiness without creating fatigue"], "A useful warm-up raises temperature, rehearses the planned movements and confirms that the body feels ready. It need not be long or exhausting. Specific practice sets are usually more valuable than unrelated high-intensity drills."),
    ("Ramp-Up Sets", "ramp_sets", GOLD, ["First set very light", "Add load in 2-4 steps", "Reduce repetitions as load rises", "Stop ramping before fatigue"], "Before a heavy compound exercise, use several submaximal practice sets. Example: 8 easy reps, 5 moderate reps, then 2-3 reps near the work weight. Accessories often need only one light rehearsal set."),
    ("Mobility for Training", "mobility", "7356A5", ["Target the position required today", "Use active control", "Avoid forcing painful end range", "Retest the exercise after the drill"], "Mobility is useful when it improves a needed training position. Choose one or two drills, perform controlled repetitions, then retest the squat, press or pull. If no useful change occurs, adjust the machine or choose a different exercise."),
    ("Stretching and Cooldown", "stretching", "7356A5", ["Slow breathing", "Mild-to-moderate stretch sensation", "No bouncing or numbness", "Use cooldown for transition, not punishment"], "Static stretching after training can support flexibility work, but it does not need to be painful. Hold comfortable positions for about 20-45 seconds and repeat as needed. Cooling down gradually is especially sensible after hard intervals."),
    ("Volume and Frequency", "volume", "2E8555", ["Train major muscle groups at least twice weekly when practical", "Begin with modest weekly sets", "Add work only when recovery remains good", "Consistency outranks complexity"], "The 2026 ACSM resistance-training update emphasizes that moving from no training to regular training produces the largest benefit. Beginners can progress with a small number of hard, repeatable sets and do not need advanced methods."),
    ("Repetition Ranges and Rest", "reps_rest", "2E8555", ["Strength often uses lower reps and longer rest", "Muscle growth can occur across a broad range", "Isolation work often uses moderate-to-high reps", "Rest until technique can repeat"], "The handbook uses broad ranges rather than a single perfect number. Heavy compound work usually needs two or more minutes. Small accessories may be ready again in 45-90 seconds. Short rest should not turn controlled strength work into accidental cardio."),
    ("Progressive Overload", "overload", GREEN, ["Add repetitions", "Add a small amount of load", "Improve range or pause quality", "Add a set only when needed"], "Progress is a trend in completed work and quality, not a requirement to add weight every session. The default method is double progression: reach the top of the rep range across all sets, then make the smallest practical load increase."),
    ("Plateaus and Deloads", "deload", GOLD, ["Check sleep, nutrition and technique", "Hold load steady for another week", "Reduce volume for 5-10 days when fatigue accumulates", "Change exercises only for a reason"], "A plateau is not one difficult workout. Look for several weeks without progress despite good attendance. A deload can reduce sets by roughly one-third to one-half while keeping movements familiar and far from failure."),
    ("Exercise Order", "exercise_order", "C67A24", ["Skill and heavy compounds first", "Secondary compounds next", "Isolation and core later", "Conditioning placement depends on the main goal"], "Place the exercise that matters most while attention and energy are high. A beginner session usually needs one lower-body compound, one push, one pull and a small amount of accessory or core work."),
    ("Choosing a Split", "splits", "C67A24", ["Full body for 2-3 days", "Upper/lower for 4 days", "Push/pull/legs for 3-6 days", "Choose by schedule and recovery"], "A split is an organizational tool. Weekly exercise selection, effort and recovery matter more than the label. Missing sessions frequently is a sign that the plan is too complex for the schedule."),
    ("Cardio Intensity and Progression", "cardio", "16889E", ["Use talk test or perceived effort", "Build minutes before intensity", "Include easy work and optional intervals", "Strength and cardio can coexist"], "HHS guidance recommends 150-300 minutes of moderate aerobic activity weekly plus muscle-strengthening activity on at least two days. Inactive adults can start with much smaller amounts and build over time."),
    ("Nutrition, Protein and Recovery", "nutrition", GREEN, ["Build meals around nutrient-dense foods", "Include protein across the day", "Hydrate before and after training", "Sleep and total energy intake shape recovery"], "Nutrition chapters in a gym manual should stay practical. Current federal guidance emphasizes an overall healthy dietary pattern. Protein supports resistance-training adaptation, but supplements are optional; food, total intake and consistent training remain foundational."),
]


def add_principles(doc: Document):
    for index, (title, bookmark, color, points, body) in enumerate(PRINCIPLES):
        if index % 2 == 0:
            start_page(doc, title, kicker="Training foundations", color=color, bookmark=bookmark)
        else:
            start_section_on_current_page(doc, title, kicker="Training foundations", color=color, bookmark=bookmark)
        add_callout(doc, "Key idea", points[0], color, "F7F9FA")
        p = doc.add_paragraph(style="Guide Body")
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(body)
        set_run(r, size=10, color=INK)
        add_bullets(doc, points)
        if bookmark == "volume":
            add_callout(doc, "Current evidence", "ACSM's 2026 update highlights consistency, individualized load and volume, and the optional nature of many advanced techniques.", "2E6FA7", "EDF3F8")
        elif bookmark == "cardio":
            add_callout(doc, "Public-health target", "Adults can work toward 150-300 minutes of moderate aerobic activity weekly and muscle strengthening on at least two days, starting smaller when inactive.", "16889E", "EAF7FA")
        elif bookmark == "nutrition":
            add_callout(doc, "Protein context", "A large meta-analysis found that protein supplementation can modestly augment resistance-training gains; benefits plateaued around a total intake of roughly 1.6 g/kg/day in the analyzed data. Individual needs vary.", GREEN, "EDF7F1")
        else:
            add_callout(doc, "Apply it", "Use the relevant guide's settings log and form cues to turn this principle into a repeatable setup.", color, "F7F9FA")
        add_label_value(doc, "Self-check", "Could another trained observer recognize the same setup and repetition from your notes?")


def guide_bookmark(g: dict) -> str:
    return f"guide_{g['no']:02d}"


def add_guide_pages(doc: Document, g: dict):
    color = CATEGORY_COLORS[g["category"]]
    bookmark = guide_bookmark(g)
    start_page(doc, f"{g['no']:02d}. {g['title']}", kicker=f"{g['category']} guide - identify", color=color, bookmark=bookmark, subtitle=f"{g['pattern']}  |  {g['level']}  |  {g['role']}")
    hero = annotated_asset(g, g["photos"][0], "_bikes" if g["no"] == 44 else "")
    add_image(doc, hero, width=6.25, alt=f"Annotated photograph of {g['title']} based on {photo_name(g['photos'][0])}")
    p = doc.add_paragraph(style="Small Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    photo_labels = ", ".join(photo_name(n) for n in g["photos"])
    p.add_run(f"Source photo mapping: {photo_labels}")
    add_callout(doc, f"Identification confidence: {g['confidence']}", g["evidence"], color, "F7F9FA")
    add_label_value(doc, "Purpose", g["purpose"], color)

    start_page(doc, f"{g['title']}: Setup and Technique", kicker=f"{g['category']} guide - perform", color=color)
    add_label_value(doc, "Movement pattern", g["pattern"], color)
    add_label_value(doc, "Primary muscles", g["primary"], color)
    add_label_value(doc, "Secondary muscles", g["secondary"], color)
    p = doc.add_paragraph(style="Heading 2")
    p.add_run("Set up")
    add_bullets(doc, g["setup"], numbered=True)
    p = doc.add_paragraph(style="Heading 2")
    p.add_run("Perform")
    if len(g["execution"]) == 1:
        position_labels = ["Movement"]
    elif len(g["execution"]) == 2:
        position_labels = ["Start position", "Movement and return"]
    elif len(g["execution"]) == 3:
        position_labels = ["Start position", "Movement", "Finish and return"]
    else:
        position_labels = ["Start position", "Movement", "Finish position", "Return"]
    execution_items = [
        f"{position_labels[min(i, len(position_labels) - 1)]}: {item}"
        for i, item in enumerate(g["execution"])
    ]
    add_bullets(doc, execution_items, numbered=True)
    add_callout(doc, "Breathing and tempo", f"Inhale during the controlled return, brace as needed, and exhale through the main effort. Suggested tempo: {g['tempo']}.", color, "F7F9FA")
    add_label_value(doc, "Form cues", "  |  ".join(g["cues"]), color)

    start_page(doc, f"{g['title']}: Programming and Troubleshooting", kicker=f"{g['category']} guide - progress", color=color)
    add_callout(doc, "Starter prescription", g["programming"], color, "F7F9FA")
    add_label_value(doc, "Workout placement", g["placement"], color)
    add_label_value(doc, "Progression", g["progression"], GREEN)
    p = doc.add_paragraph(style="Heading 2")
    p.add_run("Common mistakes and corrections")
    for mistake, correction in g["mistakes"]:
        p = doc.add_paragraph(style="Guide Body")
        r = p.add_run(mistake + ": ")
        set_run(r, size=9.5, bold=True, color=RED)
        r = p.add_run(correction)
        set_run(r, size=9.5, color=INK)
    add_callout(doc, "Safety", g["safety"], RED, "FBEDED")
    add_label_value(doc, "Alternatives", "; ".join(g["alternatives"]), color)
    if len(g["photos"]) > 1:
        paths = [plain_asset(n) for n in g["photos"][:2]]
        add_two_images(doc, paths, [f"Alternate view from {photo_name(n)} for {g['title']}" for n in g["photos"][:2]])


PROGRAM_PAGES = [
    ("Choosing the Right Program", "program_choose", ["Use full body when training 2-3 days per week.", "Use upper/lower when four consistent days are realistic.", "Use push/pull/legs when you enjoy category-based sessions and can recover.", "Start below your maximum schedule and earn more volume through consistency."]),
    ("Beginner Full Body A", "program_fba", ["Warm-up: 5-8 minutes easy cardio plus ramp-up sets.", "Leg press - 3 x 8-12.", "Plate-loaded chest press - 3 x 8-12.", "Selectorized lat pulldown - 3 x 8-12.", "Seated leg curl - 2 x 10-15.", "Exercise-mat dead bug - 2 x 6-10 per side."]),
    ("Beginner Full Body B", "program_fbb", ["Warm-up: easy cardio plus movement rehearsal.", "Smith squat or supported split squat - 3 x 8-12.", "Plate-loaded shoulder press - 3 x 8-12.", "Chest-supported row - 3 x 8-12.", "Booty Builder hip thrust - 2-3 x 8-15.", "Standing calf raise - 2 x 10-20."]),
    ("Push / Pull / Legs Overview", "program_ppl", ["Train on three nonconsecutive days at first.", "Keep most sets at 1-3 RIR.", "Use one main compound, one secondary compound and two or three accessories.", "Repeat the same version for 4-8 weeks before changing exercises."]),
    ("Push Day", "program_push", ["Chest press - 3 x 6-12.", "Incline press - 3 x 8-12.", "Shoulder press - 2-3 x 8-12.", "Pec fly - 2 x 10-20.", "Seated dip - 2 x 8-15."]),
    ("Pull Day", "program_pull", ["Lat pulldown or assisted pull-up - 3 x 6-12.", "Chest-supported row - 3 x 8-12.", "Cable row - 2 x 10-15.", "Rear-delt fly - 2 x 12-20.", "Biceps curl - 2-3 x 10-15."]),
    ("Leg Day", "program_legs", ["Leg press - 3 x 6-12.", "Hip thrust or glute trainer - 3 x 8-15.", "Seated leg curl - 3 x 8-15.", "Leg extension - 2 x 10-20.", "Standing calf raise - 3 x 10-20."]),
    ("Upper / Lower Overview", "program_ul", ["Example: Monday upper, Tuesday lower, Thursday upper, Friday lower.", "Use different rep emphasis across the two weekly exposures.", "Keep weekly sets modest until recovery is reliable.", "Cardio can follow easy sessions or occur on separate days."]),
    ("Upper Day", "program_upper", ["Flat or incline bench press - 3 x 5-10.", "Lat pulldown - 3 x 6-12.", "Shoulder press - 2-3 x 8-12.", "Seated cable row - 2-3 x 8-15.", "Pec/rear-delt fly - 2 x 12-20.", "Curl plus dip - 2 x 10-15 each."]),
    ("Lower Day", "program_lower", ["Leg press - 3 x 6-12.", "Hip thrust - 3 x 8-15.", "Leg curl - 3 x 8-15.", "Leg extension - 2 x 12-20.", "Calf raise - 3 x 10-20.", "Back extension or core bench - 2 x 10-15."]),
    ("Intermediate Four-Day Structure", "program_int", ["Upper strength emphasis.", "Lower strength emphasis.", "Upper hypertrophy emphasis.", "Lower hypertrophy emphasis.", "Use 2-4 hard sets per selected movement and manage total weekly volume."]),
    ("Intermediate Upper Session", "program_int_upper", ["Barbell bench press - 4 x 4-8.", "Plate-loaded pulldown - 4 x 6-10.", "Chest-supported row - 3 x 8-12.", "Incline press - 3 x 8-12.", "Rear-delt fly - 3 x 12-20.", "Biceps and triceps - 2-3 x 8-15."]),
    ("Intermediate Lower Session", "program_int_lower", ["Leg press - 4 x 5-10.", "Hip thrust - 4 x 6-12.", "Seated or lying leg curl - 3 x 8-15.", "Leg extension - 3 x 10-20.", "Calf raise - 4 x 8-20.", "Back extension - 2 x 10-15."]),
    ("Integrating Cardio", "program_cardio", ["Easy cardio can follow strength without major disruption.", "Place hard intervals away from demanding leg sessions when possible.", "Build weekly minutes gradually.", "Use treadmill, recumbent bike or upper-body ergometer according to joint comfort and recovery."]),
    ("Deload and Progression Week", "program_deload", ["Reduce sets by about one-third to one-half.", "Keep familiar exercises and use 3-5 RIR.", "Review settings, sleep and soreness.", "Resume progression only when technique and motivation recover."]),
]


def add_programs(doc: Document):
    for index, (title, bookmark, items) in enumerate(PROGRAM_PAGES):
        if index % 2 == 0:
            start_page(doc, title, kicker="Programs", color="C67A24", bookmark=bookmark)
        else:
            start_section_on_current_page(doc, title, kicker="Programs", color="C67A24", bookmark=bookmark)
        add_callout(doc, "Default effort", "Unless noted, finish most work sets with 1-3 technically sound repetitions still possible.", "C67A24", "FFF5E8")
        add_bullets(doc, items, numbered=True)
        add_label_value(doc, "Rest", "Compounds: generally 90-240 seconds. Accessories: generally 45-120 seconds.", "C67A24")
        add_label_value(doc, "Progression", "Reach the top of the rep range across all sets, then increase by the smallest practical amount.", GREEN)
        add_callout(doc, "Record", "Write the machine setting, load, repetitions and RIR. A repeatable setup is part of the result.", GREEN, "EDF7F1")


def add_blank_lines(doc, count=10):
    for _ in range(count):
        p = doc.add_paragraph("________________________________________________________________________________")
        p.paragraph_format.space_after = Pt(7)
        for r in p.runs:
            set_run(r, size=8, color=RULE)


def add_logs(doc: Document):
    start_page(doc, "Machine Settings Log", kicker="Track", color=GREEN, bookmark="log_settings")
    table = doc.add_table(rows=1, cols=4)
    for i, t in enumerate(("Equipment", "Seat/pad", "Load", "Notes")):
        table.rows[0].cells[i].text = t
        set_cell_shading(table.rows[0].cells[i], "E8F2EC")
    for _ in range(12):
        set_row_height(table.add_row())
    set_table_geometry(table, [2600, 1800, 1400, 3560])
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    start_page(doc, "Weekly Training Planner", kicker="Track", color=GREEN, bookmark="log_week")
    table = doc.add_table(rows=1, cols=4)
    for i, t in enumerate(("Day", "Session", "Cardio/min", "Recovery note")):
        table.rows[0].cells[i].text = t
        set_cell_shading(table.rows[0].cells[i], "E8F2EC")
    for day in ("Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"):
        row = table.add_row()
        set_row_height(row)
        cells = row.cells
        cells[0].text = day
    set_table_geometry(table, [1500, 3400, 1600, 2860])
    set_table_borders(table)
    add_callout(doc, "Weekly review", "Did the plan fit real life? What should be repeated, reduced or moved next week?", GREEN, "EDF7F1")
    add_blank_lines(doc, 4)

    for i in range(1, 5):
        start_page(doc, f"Strength Session Log {i}", kicker="Track", color=GREEN, bookmark=f"log_strength_{i}")
        table = doc.add_table(rows=1, cols=6)
        for j, t in enumerate(("Exercise", "Setting", "Load", "Sets", "Reps", "RIR")):
            table.rows[0].cells[j].text = t
            set_cell_shading(table.rows[0].cells[j], "E8F2EC")
        for _ in range(10):
            set_row_height(table.add_row())
        set_table_geometry(table, [2700, 1500, 1300, 1100, 1500, 1260])
        set_table_borders(table)
        add_label_value(doc, "Session note", "Technique, pain-free range, energy and next-session adjustment.", GREEN)
        add_blank_lines(doc, 3)

    start_page(doc, "Cardio Log", kicker="Track", color="16889E", bookmark="log_cardio")
    table = doc.add_table(rows=1, cols=6)
    for j, t in enumerate(("Date", "Mode", "Minutes", "Level", "RPE", "Notes")):
        table.rows[0].cells[j].text = t
        set_cell_shading(table.rows[0].cells[j], "EAF7FA")
    for _ in range(12):
        set_row_height(table.add_row())
    set_table_geometry(table, [1200, 1900, 1200, 1200, 1000, 2860])
    set_table_borders(table, "B9D6DC")

    start_page(doc, "Performance and Measurement Log", kicker="Track", color=GREEN, bookmark="log_performance")
    add_label_value(doc, "Choose meaningful measures", "Examples: consistent-rep load, treadmill duration, waist measurement, resting heart rate or pull-up assistance.", GREEN)
    table = doc.add_table(rows=1, cols=5)
    for j, t in enumerate(("Date", "Measure", "Result", "Conditions", "Next check")):
        table.rows[0].cells[j].text = t
        set_cell_shading(table.rows[0].cells[j], "E8F2EC")
    for _ in range(10):
        set_row_height(table.add_row())
    set_table_geometry(table, [1300, 2100, 1400, 2900, 1660])
    set_table_borders(table)

    start_page(doc, "Monthly Review", kicker="Track", color=GREEN, bookmark="log_monthly")
    for q in ("What improved?", "Which setup became more repeatable?", "Where did discomfort appear?", "What disrupted consistency?", "What is one change for next month?"):
        add_label_value(doc, q, "", GREEN)
        add_blank_lines(doc, 2)

    start_page(doc, "Safety and Equipment Checklist", kicker="Track", color=RED, bookmark="log_safety")
    add_bullets(doc, ["Selector pins intact and fully inserted", "Cables and straps free of visible damage", "Pads secure and dry", "Plate horns and collars sound", "Rack J-hooks and safeties correctly positioned", "Cardio emergency controls known", "Walkways clear", "Damaged equipment reported to staff"])
    add_callout(doc, "Do not improvise repairs", "A temporary fix can create a hidden failure. Stop using the station and notify staff.", RED, "FBEDED")
    add_blank_lines(doc, 5)


def photo_mapping() -> dict[int, list[dict]]:
    mapping = {i: [] for i in [0] + list(range(2, 52))}
    for g in GUIDES:
        for n in g["photos"]:
            mapping[n].append(g)
    return mapping


def add_indices(doc: Document):
    mapping = photo_mapping()
    nums = [0] + list(range(2, 52))
    chunks = [nums[:17], nums[17:34], nums[34:]]
    for idx, chunk in enumerate(chunks, 1):
        start_page(doc, f"Photo Index {idx} of 3", kicker="Reference", color="2E6FA7", bookmark=f"photo_index_{idx}")
        table = doc.add_table(rows=1, cols=4)
        for j, t in enumerate(("Source filename", "Guide", "Equipment identification", "Confidence")):
            table.rows[0].cells[j].text = t
            set_cell_shading(table.rows[0].cells[j], "E8EEF5")
        for n in chunk:
            gs = mapping[n]
            cells = table.add_row().cells
            cells[0].text = photo_name(n)
            cells[1].text = ", ".join(str(g["no"]) for g in gs) if gs else "Cross-reference only"
            cells[2].text = "; ".join(g["title"] for g in gs) if gs else "Plate-loaded area overview"
            cells[3].text = ", ".join(sorted(set(g["confidence"] for g in gs))) if gs else "Medium"
        set_table_geometry(table, [2300, 900, 4300, 1860])
        set_repeat_table_header(table.rows[0])

    sorted_guides = sorted(GUIDES, key=lambda g: g["title"])
    equipment_chunks = [sorted_guides[i:i + 12] for i in range(0, len(sorted_guides), 12)]
    for idx, chunk in enumerate(equipment_chunks, 1):
        start_page(doc, f"Equipment Index {idx} of {len(equipment_chunks)}", kicker="Reference", color="2E6FA7", bookmark=f"equipment_index_{idx}")
        table = doc.add_table(rows=1, cols=3)
        for j, t in enumerate(("Equipment", "Guide", "Page")):
            table.rows[0].cells[j].text = t
            set_cell_shading(table.rows[0].cells[j], "E8EEF5")
        for g in chunk:
            cells = table.add_row().cells
            cells[0].text = g["title"]
            cells[1].text = f"{g['no']:02d}"
            cells[2].text = str(PAGE_NUMBERS.get(guide_bookmark(g), ""))
        set_table_geometry(table, [6500, 1300, 1560])
        set_repeat_table_header(table.rows[0])

    glossary = [
        ("Abduction", "Movement away from the body's midline."), ("Adduction", "Movement toward the body's midline."),
        ("Brace", "Create trunk stiffness before and during a demanding repetition."), ("Cable attachment", "A handle, rope or bar connected with a carabiner."),
        ("Compound exercise", "An exercise involving more than one major joint action."), ("Concentric", "The phase in which the working muscle generally shortens."),
        ("Eccentric", "The controlled lowering phase in which the working muscle generally lengthens."), ("Failure", "The point at which another repetition cannot be completed with the chosen standard."),
        ("Hip hinge", "Bending primarily at the hips while controlling the spine."), ("Hypertrophy", "An increase in muscle size."),
        ("Isolation exercise", "An exercise focused mainly on one joint action."), ("Neutral wrist", "A wrist position close to straight rather than sharply bent."),
        ("Pivot", "The machine axis around which a lever rotates."), ("Progressive overload", "Gradually increasing training demand over time."),
        ("Range of motion", "The controlled distance traveled during a repetition."), ("RIR", "Repetitions in reserve; an estimate of technically sound reps remaining."),
        ("RPE", "Rating of perceived exertion; a subjective effort scale."), ("Selectorized", "Resistance chosen with a pin in a weight stack."),
        ("Tempo", "The planned speed of a repetition's phases."), ("Training volume", "The amount of work completed, often described with sets and repetitions."),
    ]
    gl_chunks = [glossary[:7], glossary[7:14], glossary[14:]]
    for idx, chunk in enumerate(gl_chunks, 1):
        start_page(doc, f"Glossary {idx} of 3", kicker="Reference", color="7356A5", bookmark=f"glossary_{idx}")
        for term, definition in chunk:
            add_label_value(doc, term, definition, "7356A5")

    faq_pages = [
        [
            ("Which weight should I start with?", "Choose one that allows the low end of the rep range with several clean repetitions still possible."),
            ("Do machines count as strength training?", "Yes. Machines, free weights, bands and bodyweight can all provide useful resistance."),
            ("Should I train to failure?", "Usually not. Most beginners progress well while stopping with 1-3 RIR."),
            ("What if the seat numbers are worn off?", "Record another repeatable landmark and ask staff about maintenance."),
            ("Why do left and right arms feel different?", "Use a manageable load, match range and consider unilateral work. Persistent pain or weakness warrants professional assessment."),
        ],
        [
            ("How often should I change exercises?", "Keep useful movements long enough to learn and progress them; change for goals, tolerance, access or a genuine plateau."),
            ("Is soreness required?", "No. Soreness is not a reliable measure of workout quality."),
            ("Can cardio and lifting be done together?", "Yes. Sequence them according to the main goal and manage total fatigue."),
            ("What if the handbook and placard disagree?", "Follow the current machine placard and ask staff. Machines can be replaced or configured differently."),
            ("Where is the StepMill guide?", "There is no verified StepMill in the supplied photographs, so no StepMill claim is made."),
        ],
    ]
    for idx, questions in enumerate(faq_pages, 1):
        start_page(doc, f"Frequently Asked Questions {idx} of 2", kicker="Reference", color=GOLD, bookmark=f"faq_{idx}")
        for q, a in questions:
            add_callout(doc, q, a, GOLD, "FFF8EB")

    for idx, chunk in enumerate((SOURCES[:7], SOURCES[7:]), 1):
        start_page(doc, f"References {idx} of 2", kicker="Sources", color="2E6FA7", bookmark=f"references_{idx}", subtitle="Authoritative training guidance and manufacturer identification sources")
        for citation, url in chunk:
            p = doc.add_paragraph(style="Guide Body")
            r = p.add_run(citation + " ")
            set_run(r, size=9.3, color=INK)
            add_external_link(p, "Source", url)
        add_callout(doc, "Source policy", "Manufacturer sources support equipment identification and adjustment terminology. Training recommendations are framed for healthy adults and remain general rather than individualized medical advice.", "2E6FA7", "EDF3F8")


def toc_entries():
    entries = [
        ("Publication Note and Safety Disclaimer", "disclaimer"),
        ("How to Use This Handbook", "how_to_use"),
        ("Visual Key", "visual_key"),
        ("Identification and Confidence Method", "id_method"),
        ("Photo-Mapping Overview", "photo_overview"),
    ]
    entries.extend((title, bookmark) for title, bookmark, *_ in PRINCIPLES)
    entries.extend((f"{g['no']:02d}. {g['title']}", guide_bookmark(g)) for g in GUIDES)
    entries.extend((title, bookmark) for title, bookmark, _ in PROGRAM_PAGES)
    entries.extend([
        ("Machine Settings Log", "log_settings"), ("Weekly Training Planner", "log_week"),
        ("Strength Session Logs", "log_strength_1"), ("Cardio Log", "log_cardio"),
        ("Performance and Measurement Log", "log_performance"), ("Monthly Review", "log_monthly"),
        ("Safety and Equipment Checklist", "log_safety"), ("Photo Index", "photo_index_1"),
        ("Equipment Index", "equipment_index_1"), ("Glossary", "glossary_1"),
        ("Frequently Asked Questions", "faq_1"), ("References", "references_1"),
    ])
    return entries


def insert_toc_pages(doc: Document):
    """TOC pages are added before principles after front matter.

    Page numbers are deterministic because every page is explicitly broken and
    content is constrained to one page. We first reserve four pages, then fill them
    after the document is otherwise complete by editing the stored table cells.
    """
    placeholders = []
    toc_page_count = 4
    toc_rows_per_page = 13
    for i in range(toc_page_count):
        start_page(doc, f"Contents {i+1} of {toc_page_count}", kicker="Navigate", color="2E6FA7", bookmark=f"toc_{i+1}")
        table = doc.add_table(rows=1, cols=4)
        table.rows[0].cells[0].text = "Section or guide"
        table.rows[0].cells[1].text = "Page"
        table.rows[0].cells[2].text = "Section or guide"
        table.rows[0].cells[3].text = "Page"
        for cell in table.rows[0].cells:
            set_cell_shading(cell, "E8EEF5")
        for _ in range(toc_rows_per_page):
            table.add_row()
        set_table_geometry(table, [3930, 750, 3930, 750])
        set_repeat_table_header(table.rows[0])
        placeholders.append(table)
    return placeholders


def fill_toc(tables):
    entries = toc_entries()
    rows_per_page = len(tables[0].rows) - 1
    capacity_per_page = rows_per_page * 2
    chunks = [entries[i:i + capacity_per_page] for i in range(0, len(entries), capacity_per_page)]
    if len(chunks) > len(tables):
        raise RuntimeError("TOC entry count exceeds reserved table capacity")
    for table, chunk in zip(tables, chunks):
        for row_idx in range(rows_per_page):
            for column_group, entry_idx in ((0, row_idx), (2, rows_per_page + row_idx)):
                title_cell = table.rows[row_idx + 1].cells[column_group]
                page_cell = table.rows[row_idx + 1].cells[column_group + 1]
                title_cell.text = ""
                page_cell.text = ""
                if entry_idx < len(chunk):
                    title, bookmark = chunk[entry_idx]
                    add_internal_link(title_cell.paragraphs[0], title, BOOKMARKS[bookmark])
                    page_cell.text = str(PAGE_NUMBERS[bookmark])


def build_document():
    global PAGE_COUNT, BOOKMARK_ID, BOOKMARKS, PAGE_NUMBERS
    BUILD_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    PAGE_COUNT = 0
    BOOKMARK_ID = 1
    BOOKMARKS = {}
    PAGE_NUMBERS = {}
    doc = Document()
    configure_styles(doc)
    configure_header_footer(doc)
    doc.core_properties.title = "The Complete Gym Equipment Handbook"
    doc.core_properties.subject = "Photo-mapped gym equipment reference"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.keywords = "gym equipment, strength training, cardio, setup, safety"

    make_cover(doc)
    add_front_matter(doc)
    toc_tables = insert_toc_pages(doc)
    add_principles(doc)
    for g in GUIDES:
        add_guide_pages(doc, g)
    add_programs(doc)
    add_logs(doc)
    add_indices(doc)
    fill_toc(toc_tables)

    if PAGE_COUNT != 187:
        raise RuntimeError(f"Planned page block count is {PAGE_COUNT}, expected 187")

    doc.save(OUTPUT_DOCX)
    print(f"DOCX={OUTPUT_DOCX}")
    print(f"PLANNED_PAGES={PAGE_COUNT}")
    print(f"GUIDES={len(GUIDES)}")
    print(f"PHOTOS_MAPPED={len(photo_mapping())}")


if __name__ == "__main__":
    build_document()
